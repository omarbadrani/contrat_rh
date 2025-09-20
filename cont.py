import logging
import shutil
import subprocess
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import datetime
import re
import pymysql
import pytz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches
import ttkbootstrap as ttk
from plyer import notification
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
from ttkbootstrap.tableview import Tableview
from ttkbootstrap.validation import add_regex_validation
from ttkbootstrap.widgets import DateEntry
from docx.shared import RGBColor
import zipfile
from io import BytesIO
import os
import winsound
import threading
import tempfile
import time
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from ttkbootstrap.constants import *

import logging
import logging.handlers
import os
import sys


# Configuration du logging
def setup_logging():
    # Créer le dossier logs s'il n'existe pas
    log_dir = os.path.join(os.path.dirname(__file__), 'logs')
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)

    # Fichier de log principal
    log_file = os.path.join(log_dir, 'contrat_app.log')

    # Configuration du logging
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.handlers.RotatingFileHandler(
                log_file, maxBytes=5 * 1024 * 1024, backupCount=5
            ),
            logging.StreamHandler(sys.stdout)
        ]
    )

    # Logger spécifique pour l'application
    logger = logging.getLogger('ContratApp')
    return logger


# Initialiser le logger
logger = setup_logging()
import logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
from dateutil.relativedelta import relativedelta
import re

class ContratApplication:

    def __init__(self, root):



        self.CDD_MASCULIN = """
        طبقًا لأحكام الفصل 6-4 الجديد من مجلة الشغل

        تمهيد:حيث أن شركة أمبار منيف مختصة في صناعة الأحذية للتصدير"Sous Traitance" لفائدة حرفاء بالخارج تفوق في بعض الأحيان حجم قدرتها الإنتاجية العادية، ولما كانت هذه الطلبيات محدودة وغير منتظمة وغير مضمونة الاستمرارية فإن شركة أمبار منيف بحاجة إلى انتداب أجراء لمدة معينة لتلبية حاجياتها الظرفية نظرا للزيادة الغير عادية في حجم العمل المتأتية من ارتفاع طلبيات الحرفاء: Décathlon, Imac , Ricker et Cleon  . يعتبر الطرفان هذا التمهيد جزء من العقد ويلتزمان به كسائر فصوله.
        الحمد لله،
        -  عملا بأحكام مجلة الشغل ومقتضيات الاتفاقية  المشتركة القطاعية لصناعة الأحذية وتوابعها،
        - وبناء على طلب   {{Titre}} {{Prénom}} {{Nom}}  وتصريحه بأنه حر من كل التزام وغير مرتبط بعلاقة شغلية مع أي  مؤجر كان،
        تم الاتفاق والتراضي والتعاقد  بين الممضين أسفله
        أولا : شركة أمبار منيف، شركة خفية الاسم، مقرها الاجتماعي طريق قابس 3060 المحرس، ممثلة في شخص السيد ياسين المنيف بصفته مديرا عاما، المؤجر، من جهة
        ثانيا :  {{Titre}} {{Prénom}} {{Nom}} ، تاريخ الولادة {{DNAIS}}. مكانها {{LNAIS}} عنوان  {{Ville}} صاحب بطاقة التعريف القومية عدد  {{NCIN}} الصادرة بتاريخ {{DCIN}} بـ {{LCIN}}، بصفته متعاقد لمدة معينة، من جهة أخرى
        الفصل الأول : انتدبت شركة أمبار منيف بمقتضى هذا العقد {{Titre}} {{Prénom}} {{Nom}}  ليقع تشغيله بصفة {{Poste}} وذلك بمصنعها الكائن بالعنوان المذكور أو بجميع المصالح التابعة لها.
        الفصل الثاني : مدة العمل وطبيعته : اتفق الطرفان على أن تبدأ العلاقة الشغلية بتاريخ {{DPERIODE}} وتنتهي  بتاريخ  {{FPERIODE}}، وذلك لمدة محدودة قدرها {{DUREE}}.
        الفصل الثالث : يتقاضى {{Titre}} {{Prénom}} {{Nom}} مقابل انجاز العمل موضوع هذا العقد أجرا خاما خاضعا للتنقيصات الاجتماعية والجنائية تفصيله كالآتي :
          أجر أساسي: {{SBASE}} دينار في  {{MPAIE}}
          منح مختلفة: {{PRIME}} دينارًا عن الشهر الكامل
        الفصل الرابع : يعترف الأجير أنه اطلع على تراتيب العمل السارية داخل المؤسسة وتعهد باحترامها بصفتها جزءا لا يتجزأ عن هذا العقد بما في ذلك أوقات العمل، كما يتعهد بالعناية بعملها والمحافظة على ممتلكات المؤسسة من مواد أولية وأدوات وآلات وغيرها، كما أن عليها قبول أي إحالة إلى مركز عمل أو مصلحة يقترحها عليها المؤجر.
        الفصل الخامس : ينتهي هذا العقد بانتهاء أجله المذكور إعلاه بدون سابق إعلام ويصبح الطرفان في حل من العلاقة الشغلية التي تربطهما. كما ينتهي هذا العقد باتفاق الطرفين أو عند انتهاء الأشغال موضوع هذا العقد أو عند قيام الأجير بهفوة فادحة أو ضعف في مؤهلاته أو تدني إنتاجه.
        الفصل السادس : يتعهد  {{Titre}} {{Prénom}} {{Nom}} باحترام سمعة المؤسسة والإمتناع عن إفشاء الأسرار المهنية والصناعية مما ينجر عنه الإضرار بمصالح المؤسسة حتى بعد انتهاء العقد. ويتحمل الأجير مسؤولية كل ضرر ينجر عن مخالفة هذا الشرط ويتعهد بأداء الغرامات اللازمة لجبره بدون مساس بالتتبعات الجزائية طبقا للقانون الجاري به العمل.
        الفصل السابع : يعترف الطرفان أنهما اطلعا على هذا العقد وأمضيا في حالة جواز ومعرفة، ويعينان مقر مخابرتهما بعنوانهما المذكور أعلاه كما يلتزم الأجير بإعلام مؤجره كتابيا بكل تغيير يطرأ عن البيانات الشخصية التي أدل بها بما في ذلك عنوانه ومقر سكناه في ظرف 48 ساعة من تاريخ التغيير.
        حرر بالمحرس في نظيرين بتاريخ {{DATE_CONTRAT}}
          	إمضاء المؤجر وختمه                                                            	 	     إمضاء الأجير معرف ب
        """

        self.CDD_FEMININ = """
        طبقًا لأحكام الفصل 6-4 الجديد من مجلة الشغل 

        تمهيد:حيث أن شركة أمبار منيف مختصة في صناعة الأحذية للتصدير"Sous Traitance" لفائدة حرفاء بالخارج تفوق في بعض الأحيان حجم قدرتها الإنتاجية العادية، ولما كانت هذه الطلبيات محدودة وغير منتظمة وغير مضمونة الاستمرارية فإن شركة أمبار منيف بحاجة إلى انتداب أجراء لمدة معينة لتلبية حاجياتها الظرفية نظرا للزيادة الغير عادية في حجم العمل المتأتية من ارتفاع طلبيات الحرفاء: Décathlon, Imac , Ricker et Cleon  . يعتبر الطرفان هذا التمهيد جزء من العقد ويلتزمان به كسائر فصوله.
        الحمد لله،
        -  عملا بأحكام مجلة الشغل ومقتضيات الاتفاقية  المشتركة القطاعية لصناعة الأحذية وتوابعها،
        - وبناء على طلب   {{Titre}} {{Prénom}} {{Nom}}  وتصريحها بأنها حرة من كل التزام وغير مرتبطة بعلاقة شغلية مع أي  مؤجر كان،
        تم الاتفاق والتراضي والتعاقد  بين الممضين أسفله
        أولا : شركة أمبار منيف، شركة خفية الاسم، مقرها الاجتماعي طريق قابس 3060 المحرس، ممثلة في شخص السيد ياسين المنيف بصفته مديرا عاما، المؤجر، من جهة
        ثانيا :  {{Titre}} {{Prénom}} {{Nom}} ، تاريخ الولادة {{DNAIS}}. مكانها {{LNAIS}} عنوانها {{Ville}} صاحبة بطاقة التعريف القومية عدد  {{NCIN}} الصادرة بتاريخ {{DCIN}} بـ {{LCIN}}، بصفتها {{Poste}} متعاقدة لمدة معينة، من جهة أخرى
        الفصل الأول : انتدبت شركة أمبار منيف بمقتضى هذا العقد {{Titre}} {{Prénom}} {{Nom}}  ليقع تشغيلها بصفة  {{Poste}} وذلك بمصنعها الكائن بالعنوان المذكور أو بجميع المصالح التابعة لها.
        الفصل الثاني : مدة العمل وطبيعته : اتفق الطرفان على أن تبدأ العلاقة الشغلية بتاريخ {{DPERIODE}} وتنتهي  بتاريخ  {{FPERIODE}}، وذلك لمدة محدودة قدرها {{DUREE}}.
        الفصل الثالث : تتقاضى {{Titre}} {{Prénom}} {{Nom}} مقابل انجاز العمل موضوع هذا العقد أجرا خاما خاضعا للتنقيصات الاجتماعية والجنائية تفصيله كالآتي :
          أجر أساسي: {{SBASE}} دينار في {{MPAIE}}
          منح مختلفة: {{PRIME}} دينارًا عن الشهر الكامل
        الفصل الرابع : تعترف الأجيرة أنها اطلعت على تراتيب العمل السارية داخل المؤسسة وتعهدت باحترامها بصفتها جزءا لا يتجزأ عن هذا العقد بما في ذلك أوقات العمل، كما تتعهد بالعناية بعملها والمحافظة على ممتلكات المؤسسة من مواد أولية وأدوات وآلات وغيرها، كما أن عليها قبول أي إحالة إلى مركز عمل أو مصلحة يقترحها عليها المؤجر.
        الفصل الخامس : ينتهي هذا العقد بانتهاء أجله المذكور إعلاه بدون سابق إعلام ويصبح الطرفان في حل من العلاقة الشغلية التي تربطهما. كما ينتهي هذا العقد باتفاق الطرفين أو عند انتهاء الأشغال موضوع هذا العقد أو عند قيام الأجيرة بهفوة فادحة أو ضعف في مؤهلاتها أو تدني إنتاجها.
        الفصل السادس : تتعهد  {{Titre}} {{Prénom}} {{Nom}} باحترام سمعة المؤسسة والإمتناع عن إفشاء الأسرار المهنية والصناعية مما ينجر عنه الإضرار بمصالح المؤسسة حتى بعد انتهاء العقد. وتتحمل الأجيرة مسؤولية كل ضرر ينجر عن مخالفة هذا الشرط وتتعهد بأداء الغرامات اللازمة لجبره بدون مساس بالتتبعات الجزائية طبقا للقانون الجاري به العمل.
        الفصل السابع : يعترف الطرفان أنهما اطلعا على هذا العقد وأمضيا في حالة جواز ومعرفة، ويعينان مقر مخابرتهما بعنوانهما المذكور أعلاه كما تلتزم الأجيرة بإعلام مؤجرها كتابيا بكل تغيير يطرأ عن البيانات الشخصية التي أدلت بها بما في ذلك عنوانها ومقر سكناها في ظرف 48 ساعة من تاريخ التغيير.
        حرر بالمحرس في نظيرين بتاريخ {{DATE_CONTRAT}}
          	إمضاء المؤجر وختمه                                                            	 	     إمضاء الأجيرة معرف ب
        """

        self.CDI_MASCULIN = """
        الحمد لله،
        -  عملا بأحكام مجلة الشغل ومقتضيات الاتفاقية المشتركة القطاعية لصناعة الأحذية وتوابعها،
        وبناء على طلب  {{Titre}} {{Prénom}} {{Nom}} وتصريحه بأنه حر من كل التزام وغير مرتبط بعلاقة شغليه مع أي مؤجر كان،
        تم الاتفاق والتراضي والتعاقد بين الممضين أسفله
        أولا: شركة أمبار منيف، شركة خفية الاسم، مقرها الاجتماعي طريق قابس 3060 المحرس، ممثلة في شخص السيد ياسين المنيف بصفته مديرا عاما، المؤجر، من جهة
        ثانيا : السيد  {{Titre}} {{Prénom}} {{Nom}} تاريخ الولادة {{DNAIS}}. مكانها {{LNAIS}} عنوانه {{Ville}} صاحب بطاقة التعريف القومية عدد {{NCIN}} الصادرة بتاريخ {{DCIN}} بتونس بصفته أجير متعاقد لمدة غير معينة، من جهة أخرى
        الفصل الأول : انتدبت شركة أمبار منيف بمقتضى هذا العقد  {{Titre}} {{Prénom}} {{Nom}} ليقع تشغيله بصفة {{Poste}} وذلك بمصنعها الكائن بالعنوان المذكور أو بجميع المصالح التابعة لها.
        الفصل الثاني : مدة العمل وطبيعته : اتفق الطرفان على أن تبدأ العلاقة الشغلية لمدة غير محددة من تاريخ {{DPERIODE}} .
         يخضع العامل المذكور أعلاه الى فترة تجربة مدتها ستة أشهر قابلة للتجديد مرة واحدة ولنفس المدة ويمكن لاحد طرفي العقد انهاء العمل به قبل انقضاء فترة التجربة بعد اعلام الطرف الأخر باي وسيلة تترك اثرا كتابيا وذلك قبل خمسة عشر يوما من انهاء فترة التجربة.
        الفصل الثالث: يتقاضى  {{Titre}} {{Prénom}} {{Nom}} مقابل انجاز العمل موضوع هذا العقد أجرا خاما خاضعا للتنقيصات الاجتماعية والجنائية تفصيله كالآتي:
        أجر أســـاسي {{SBASE}}   دينار  في الساعة
        منـــح مختلفة {{PRIME}} دينارا  عن الشهر الكامل
        الفصل الرابع: يعترف الأجير أنه اطلع على تراتيب العمل السارية داخل المؤسسة وتعهد باحترامها بصفتها جزءا لا يتجزأ عن هذا العقد بما في ذلك أوقات العمل، كما يتعهد بالعناية بعمله والمحافظة على ممتلكات المؤسسة من مواد أولية وأدوات وآلات وغيرها، كما أن عليه قبول أي إحالة إلى مركز عمل أو مصلحة يقترحها عليه المؤجر. 
        الفصل السادس : يتعهد السيد  {{Titre}} {{Prénom}} {{Nom}} باحترام سمعة المؤسسة والامتناع عن إفشاء الأسرار المهنية والصناعية مما ينجر عنه الإضرار بمصالح المؤسسة حتى بعد انتهاء العقد. ويتحمل الأجير مسؤولية كل ضرر ينجر عن مخالفة هذا الشرط ويتعهد بأداء الغرامات اللازمة لجبره بدون مساس بالتتبعات الجزائية طبقا للقانون الجاري به العمل.
        الفصل السابع: يعترف الطرفان أنهما اطلعا على هذا العقد وأمضيا في حالة جواز ومعرفة، ويعينان مقر مخابرتهما بعنوانهما المذكور أعلاه كما تلتزم الأجيرة بإعلام مؤجرها كتابيا بكل تغيير يطرأ عن البيانات الشخصية التي أدلت بها بما في ذلك عنوانها ومقر سكناها في ظرف 48 ساعة من تاريخ التغيير.
        حرر بالمحرس في نظيرين بتاريخ {{DATE_CONTRAT}}
              إمضاء المؤجر وختمه                     			                    إمضاء الأجير معرف به
        """

        self.CDI_FEMININ = """
        الحمد لله،
        -  عملا بأحكام مجلة الشغل ومقتضيات الاتفاقية المشتركة القطاعية لصناعة الأحذية وتوابعها،
        وبناء على طلب  {{Titre}} {{Prénom}} {{Nom}} وتصريحها بأنها حرة من كل التزام وغير مرتبطة بعلاقة شغليه مع أي مؤجر كان،
        تم الاتفاق والتراضي والتعاقد بين الممضين أسفله
        أولا: شركة أمبار منيف، شركة خفية الاسم، مقرها الاجتماعي طريق قابس 3060 المحرس، ممثلة في شخص السيد ياسين المنيف بصفته مديرا عاما، المؤجر، من جهة
        ثانيا : الآنسة {{Titre}} {{Prénom}} {{Nom}} تاريخ الولادة {{DNAIS}}. مكانها {{LNAIS}} عنوانها {{Ville}} صاحبة بطاقة التعريف القومية عدد {{NCIN}} الصادرة بتاريخ {{DCIN}} بتونس بصفتها أجيرة متعاقدة لمدة غير معينة، من جهة أخرى
        الفصل الأول : انتدبت شركة أمبار منيف بمقتضى هذا العقد  {{Titre}} {{Prénom}} {{Nom}} ليقع تشغيلها بصفة {{Poste}} وذلك بمصنعها الكائن بالعنوان المذكور أو بجميع المصالح التابعة لها.
        الفصل الثاني : مدة العمل وطبيعته : اتفق الطرفان على أن تبدأ العلاقة الشغلية لمدة غير محددة من تاريخ {{DPERIODE}} .
         تخضع العاملة المذكورة أعلاه الى فترة تجربة مدتها ستة أشهر قابلة للتجديد مرة واحدة ولنفس المدة ويمكن لاحد طرفي العقد انهاء العمل به قبل انقضاء فترة التجربة بعد اعلام الطرف الأخر باي وسيلة تترك اثرا كتابيا وذلك قبل خمسة عشر يوما من انهاء فترة التجربة.
        الفصل الثالث: تتقاضى  {{Titre}} {{Prénom}} {{Nom}} مقابل انجاز العمل موضوع هذا العقد أجرا خاما خاضعا للتنقيصات الاجتماعية والجنائية تفصيله كالآتي:
        أجر أســـاسي {{SBASE}}   دينار  في الساعة
        منـــح مختلفة {{PRIME}} دينارا  عن الشهر الكامل
        الفصل الرابع: تعترف الأجيرة أنها اطلعت على تراتيب العمل السارية داخل المؤسسة وتعهدت باحترامها بصفتها جزءا لا يتجزأ عن هذا العقد بما في ذلك أوقات العمل، كما تتعهد بالعناية بعملها والمحافظة على ممتلكات المؤسسة من مواد أولية وأدوات وآلات وغيرها، كما أن عليها قبول أي إحالة إلى مركز عمل أو مصلحة يقترحها عليها المؤجر. 
        الفصل السادس : تتعهد الآنسة {{Titre}} {{Prénom}} {{Nom}} باحترام سمعة المؤسسة والامتناع عن إفشاء الأسرار المهنية والصناعية مما ينجر عنه الإضرار بمصالح المؤسسة حتى بعد انتهاء العقد. وتتحمل الأجيرة مسؤولية كل ضرر ينجر عن مخالفة هذا الشرط وتتعهد بأداء الغرامات اللازمة لجبره بدون مساس بالتتبعات الجزائية طبقا للقانون الجاري به العمل.
        الفصل السابع: يعترف الطرفان أنهما اطلعا على هذا العقد وأمضيا في حالة جواز ومعرفة، ويعينان مقر مخابرتهما بعنوانهما المذكور أعلاه كما تلتزم الأجيرة بإعلام مؤجرها كتابيا بكل تغيير يطرأ عن البيانات الشخصية التي أدلت بها بما في ذلك عنوانها ومقر سكناها في ظرف 48 ساعة من تاريخ التغيير.
        حرر بالمحرس في نظيرين بتاريخ {{DATE_CONTRAT}}
              إمضاء المؤجر وختمه                     			                    إمضاء الأجيرة معرف بها
        """
        logger.info("Initialisation de l'application ContratApplication")

        self.root = root
        self.root.title("Gestion des Contrats - Imbert Mnif")
        self.root.state('zoomed')
        self.style = ttk.Style(theme='flatly')

        # Configuration de la base de données (CORRIGÉ)
        self.db_config = {
            "host": "192.168.1.210",
            "user": "omar",
            "password": "1234",
            "database": "rh1",
            "charset": "utf8mb4"
            # Suppression de 'collation' qui n'est pas supporté
        }

        try:
            self.conn = pymysql.connect(**self.db_config)
            self.create_database()
        except pymysql.MySQLError as e:
            # Correction de l'affichage d'erreur
            error_msg = f"Erreur de connexion: {str(e)}"
            print(error_msg)
            Messagebox.showerror("Erreur de connexion", error_msg)
            self.root.quit()
            return
        except Exception as e:
            error_msg = f"Erreur inattendue: {str(e)}"
            print(error_msg)
            Messagebox.showerror("Erreur", error_msg)
            self.root.quit()
            return



        self.style.configure("Treeview", rowheight=30, font=('Segoe UI', 10))
        self.style.configure("Treeview.Heading", font=('Segoe UI', 11, 'bold'))
        self.last_contract_check = None
        self.alert_timer = None
        self.check_interval = 4* 3000  # 1 minute en millisecondes
        self.alerted_contracts = {}  # Stocker les contrats alertés avec date de fin et timestamp
        self.contract_hash = {}
        self.progress_dialog = None
        self.progress_bar = None
        self.progress_label = None
        self.sound_file = os.path.normpath(r"D:\UIAlert_Notification lasolisa 4 (ID 2066)_LS.wav")
        if not os.path.exists(self.sound_file):
            Messagebox.show_warning(
                f"Le fichier sonore d'alerte n'a pas été trouvé à l'emplacement :\n{self.sound_file}\n\nLes alertes seront silencieuses.",
                "Avertissement",
                parent=root
            )
        self.entries = {}
        self.contract_entries = {}
        self.db_config = {
            "host": "192.168.1.210",
            "user": "omar",
            "password": "1234",
            "database": "rh1",
            "charset": "utf8mb4"
            # Supprimez 'collation' car il n'est pas supporté par pymysql.connect()
        }

        try:
            self.conn = pymysql.connect(**self.db_config)
            self.create_database()
        except pymysql.MySQLError as e:
            print(f"Erreur de connexion: {str(e)}", "Erreur")
            self.root.quit()
            return

        self.variables = {
            "genre": tk.StringVar(value="féminin"),
            "matricule": tk.StringVar(),
            "contract_type": tk.StringVar(value="CDD"),
            "salary_type": tk.StringVar(value="في الساعة")
        }
        self.current_employee = None
        self.logo_path = r"D:\imbertlogo.png"
        self.undo_stack = []
        self.column_definitions = [
            {"text": "Matricule", "stretch": False, "width": 100},
            {"text": "Nom", "stretch": True, "width": 150},
            {"text": "Prénom", "stretch": True, "width": 150},
            {"text": "Genre", "stretch": False, "width": 80},
            {"text": "Date Naissance", "stretch": True, "width": 120},
            {"text": "Lieu Naissance", "stretch": True, "width": 150},
            {"text": "Adresse", "stretch": True, "width": 200},
            {"text": "Ville", "stretch": True, "width": 100},
            {"text": "CIN", "stretch": True, "width": 100},
            {"text": "Date CIN", "stretch": True, "width": 120},
            {"text": "Lieu CIN", "stretch": True, "width": 150},
            {"text": "Poste", "stretch": True, "width": 150},
            {"text": "Email", "stretch": True, "width": 200},
            {"text": "Téléphone", "stretch": True, "width": 120},
            {"text": "Type Contrat", "stretch": True, "width": 100},
            {"text": "Date Début", "stretch": True, "width": 120},
            {"text": "Date Fin", "stretch": True, "width": 120},
            {"text": "Salaire Base", "stretch": True, "width": 100},
            {"text": "Prime", "stretch": True, "width": 100},
            {"text": "Type Salaire", "stretch": True, "width": 100},
            {"text": "Atelier", "stretch": True, "width": 120},
            {"text": "Nb Échéances", "stretch": False, "width": 100}

        ]
        self.alert_column_definitions = [
            {"text": "Matricule", "stretch": False, "width": 100},
            {"text": "Nom", "stretch": True, "width": 150},
            {"text": "Prénom", "stretch": True, "width": 150},
            {"text": "Date Fin", "stretch": True, "width": 120},
            {"text": "Jours Restants", "stretch": True, "width": 120},
            {"text": "Atelier", "stretch": True, "width": 120},  # Nouvelle colonne
            {"text": "Nb Échéances", "stretch": False, "width": 100} , # Nouvelle colonne
            {"text": "Type Contrat", "stretch": True, "width": 100},
            {"text": "Fperiode", "stretch": True, "width": 120}  # Nouvelle colonne
        ]

        self.setup_ui()
        self.load_data()
        self.background_alert_service()

    def debug_info(self, message):
        """Affiche les informations de débogage"""
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        debug_message = f"[DEBUG {timestamp}] {message}"
        print(debug_message)
        # Vous pouvez aussi logger dans un fichier si nécessaire

    def validate_database_connection(self):
        """Valide la connexion à la base de données"""
        try:
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT 1")
                self.debug_info("Connexion à la base de données validée")
                return True
        except Exception as e:
            self.debug_info(f"Erreur de connexion à la base: {str(e)}")
            print(f"Erreur de connexion à la base de données: {str(e)}", "Erreur")
            return False

    def safe_execute_sql(self, query, params=None):
        """Exécute une requête SQL avec gestion d'erreurs"""
        try:
            with self.conn.cursor() as cursor:
                if params:
                    cursor.execute(query, params)
                else:
                    cursor.execute(query)
                return cursor
        except Exception as e:
            self.debug_info(f"Erreur SQL: {query}, params: {params}, erreur: {str(e)}")
            raise



    def show_last_contract(self):
        if not hasattr(self, 'current_employee') or not self.current_employee:
            Messagebox.show_warning("Veuillez d'abord sélectionner un employé", "Attention")
            return

        try:
            with self.conn.cursor() as cursor:
                # Récupérer uniquement le texte du contrat
                cursor.execute('''
                               SELECT texte_contrat
                               FROM contrats
                               WHERE matricule = %s
                               ORDER BY date_creation DESC LIMIT 1
                               ''', (self.current_employee['matricule'],))

                result = cursor.fetchone()

                if result and result[0]:
                    self.display_contract(result[0])
                else:
                    # Générer un nouveau contrat si aucun n'existe
                    contrat_text = self.generate_contract_text(self.current_employee)
                    self.display_contract(contrat_text)

                    # Sauvegarder le nouveau contrat
                    cursor.execute('''
                                   INSERT INTO contrats (matricule, texte_contrat, date_creation)
                                   VALUES (%s, %s, NOW())
                                   ''', (self.current_employee['matricule'], contrat_text))
                    self.conn.commit()

        except Exception as e:
            Messagebox.showerror("Erreur", f"Impossible d'afficher le contrat: {str(e)}")

    def display_contract(self, contract_text):
        """Affiche le contrat dans l'interface"""
        self.contract_text.config(state=tk.NORMAL)
        self.contract_text.delete(1.0, tk.END)
        self.contract_text.insert(tk.END, contract_text, 'rtl')
        self.contract_text.config(state=tk.DISABLED)
        self.notebook.select(2)  # Affiche l'onglet Contrat

    def create_database(self):
        cursor = self.conn.cursor()
        cursor.execute(f'''
            CREATE TABLE IF NOT EXISTS employees (
                matricule VARCHAR(50) PRIMARY KEY,
                nom VARCHAR(100) NOT NULL,
                prenom VARCHAR(100) NOT NULL,
                genre VARCHAR(20) NOT NULL,
                date_naissance VARCHAR(10),
                lieu_naissance VARCHAR(100),
                adresse VARCHAR(200),
                ville VARCHAR(100) DEFAULT 'المحرس',
                cin VARCHAR(20),
                date_cin VARCHAR(10),
                lieu_cin VARCHAR(100) DEFAULT 'تونس',
                poste VARCHAR(100),
                email VARCHAR(100),
                telephone VARCHAR(20)
            ) CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci
        ''')
        cursor.execute(f'''
            CREATE TABLE IF NOT EXISTS contrats (
                id INT AUTO_INCREMENT PRIMARY KEY,
                matricule VARCHAR(50),
                type_contrat VARCHAR(10),
                date_creation VARCHAR(20),
                texte_contrat TEXT,
                FOREIGN KEY (matricule) REFERENCES employees(matricule)
            ) CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci
        ''')
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_matricule ON employees(matricule)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_contrat_matricule ON contrats(matricule)")
        self.conn.commit()

    def setup_ui(self):
        main_panel = ttk.Frame(self.root)
        main_panel.pack(fill=BOTH, expand=True, padx=15, pady=15)
        self.notebook = ttk.Notebook(main_panel, bootstyle=PRIMARY)
        self.notebook.pack(fill=BOTH, expand=True)

        self.create_employee_tab(main_panel)
        self.create_search_tab(main_panel)
        self.create_contract_tab(main_panel)
        self.create_list_tab(main_panel)
        self.create_alerts_tab(main_panel)

        status_frame = ttk.Frame(main_panel, bootstyle=INFO)
        status_frame.pack(fill=X, pady=(10, 0))

        self.status_var = tk.StringVar(value="Prêt")
        ttk.Label(status_frame, textvariable=self.status_var, bootstyle=(INFO, INVERSE),
                  font=('Segoe UI', 10)).pack(side=LEFT, padx=10)

        ttk.Button(status_frame, text="Aide", command=self.show_help,
                   bootstyle=(INFO, OUTLINE)).pack(side=RIGHT, padx=5)
        ttk.Button(status_frame, text="À propos", command=self.show_about,
                   bootstyle=(INFO, OUTLINE)).pack(side=RIGHT, padx=5)

        self.sound_enabled = True
        ttk.Checkbutton(
            status_frame,
            text="Alertes sonores",
            variable=tk.BooleanVar(value=self.sound_enabled),
            command=lambda: setattr(self, 'sound_enabled', not self.sound_enabled),
            bootstyle="round-toggle"
        ).pack(side=RIGHT, padx=5)

    def create_employee_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Nouvel Employé")

        canvas = tk.Canvas(frame, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview, bootstyle=PRIMARY)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        fields = [
            ("Matricule*", "matricule", r'^\w+$', ttk.Entry),
            ("Nom*", "nom", None, ttk.Entry),
            ("Prénom*", "prenom", None, ttk.Entry),
            ("Date Naissance (JJ/MM/AAAA)", "date_naissance", r'^\d{2}/\d{4}/\d{2}$', DateEntry),
            ("Lieu Naissance", "lieu_naissance", None, ttk.Entry),
            ("Adresse", "adresse", None, ttk.Entry),
            ("Ville", "ville", None, ttk.Entry),
            ("Code Postal", "code_postal", r'^\d{4}$', ttk.Entry),
            ("CIN", "cin", r'^\d{8}$', ttk.Entry),
            ("Date CIN (JJ/MM/AAAA)", "date_cin", r'^\d{4}/\d{2}/\d{2}$', DateEntry),
            ("Lieu CIN", "lieu_cin", None, ttk.Entry),
            ("Poste", "poste", None, ttk.Entry),
            ("Email", "email", r'^[^@]+@[^@]+\.[^@]+$', ttk.Entry),
            ("Téléphone", "telephone", r'^\+?\d{10,12}$', ttk.Entry),
            ("Date Embauche (JJ/MM/AAAA)", "date_embauche", r'^\d{4}/\d{2}/\d{2}$', DateEntry),
            ("Dcon", "dcon", None, ttk.Entry),
            ("Durée", "duree", None, ttk.Entry),
            ("Atelier", "atelier", None, ttk.Entry),
            ("Nb Échéances", "nbre_eche", r'^\d+$', ttk.Entry),
            ("Fperiode", "fperiode", None, ttk.Entry),
            ("Degré Polyvalence", "degre_polyvalence", None, ttk.Entry)
        ]

        self.entries = {}
        form_frame = ttk.LabelFrame(scrollable_frame, text="Informations Employé", bootstyle=PRIMARY)
        form_frame.pack(fill=tk.BOTH, padx=10, pady=10, expand=True)

        for i, (label, field, regex, widget_type) in enumerate(fields):
            ttk.Label(form_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=tk.E)
            entry = widget_type(form_frame, bootstyle="primary") if widget_type != DateEntry else DateEntry(form_frame,
                                                                                                            bootstyle="primary",
                                                                                                            dateformat="%d/%m/%Y")
            if widget_type == DateEntry:
                entry.entry.configure(justify="center")
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=tk.EW)
            if regex and widget_type == ttk.Entry:
                add_regex_validation(entry, regex)
                entry.bind("<KeyRelease>", lambda e, f=field: self.validate_field(e.widget, f))
            self.entries[field] = entry

        ttk.Label(form_frame, text="Genre*", font=('Segoe UI', 10)).grid(row=len(fields), column=0, sticky=tk.E, padx=5,
                                                                         pady=5)
        genre_frame = ttk.Frame(form_frame)
        genre_frame.grid(row=len(fields), column=1, sticky=tk.W)
        # Dans create_employee_tab(), remplacez:
        ttk.Radiobutton(genre_frame, text="السيدة", variable=self.variables["genre"], value="السيدة",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(genre_frame, text="الانسة", variable=self.variables["genre"], value="الانسة",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(genre_frame, text="السيد", variable=self.variables["genre"], value="السيد",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)
        self.entries['ville'].insert(0, "المحرس")
        self.entries['lieu_cin'].insert(0, "تونس")

        contract_frame = ttk.LabelFrame(scrollable_frame, text="Détails du Contrat", bootstyle=PRIMARY)
        contract_frame.pack(fill=tk.BOTH, padx=10, pady=10, expand=True)

        ttk.Label(contract_frame, text="Type de Contrat*", font=('Segoe UI', 10)).grid(row=0, column=0, padx=5, pady=5,
                                                                                       sticky=tk.E)
        ttk.Radiobutton(contract_frame, text="CDD", variable=self.variables["contract_type"], value="CDD",
                        bootstyle="primary-toolbutton", command=self.toggle_date_fin).grid(row=0, column=1, sticky=tk.W)
        ttk.Radiobutton(contract_frame, text="CDI", variable=self.variables["contract_type"], value="CDI",
                        bootstyle="primary-toolbutton", command=self.toggle_date_fin).grid(row=0, column=2, sticky=tk.W)

        contract_fields = [
            ("Date Début (JJ/MM/AAAA)*", "date_debut", r'^\d{4}/\d{2}/\d{2}$', DateEntry),
            ("Date Fin (JJ/MM/AAAA)", "date_fin", r'^\d{4}/\d{2}/\d{2}$', DateEntry),
            ("Salaire Base*", "salaire", r'^\d+(\.\d{1,2})?$', ttk.Entry),
            ("Prime*", "prime", r'^\d+(\.\d{1,2})?$', ttk.Entry)
        ]

        self.contract_entries = {}
        for i, (label, field, regex, widget_type) in enumerate(contract_fields):
            ttk.Label(contract_frame, text=label, font=('Segoe UI', 10)).grid(row=i + 1, column=0, padx=5, pady=5,
                                                                              sticky=tk.E)
            entry = widget_type(contract_frame, bootstyle="primary") if widget_type != DateEntry else DateEntry(
                contract_frame, bootstyle="primary", dateformat="%d/%m/%Y")
            if widget_type == DateEntry:
                entry.entry.configure(justify="center")
                if field == "date_fin" and self.variables["contract_type"].get() == "CDI":
                    entry.entry.config(state=tk.DISABLED)
            entry.grid(row=i + 1, column=1, padx=5, pady=5, sticky=tk.EW)
            if regex and widget_type == ttk.Entry:
                add_regex_validation(entry, regex)
                entry.bind("<KeyRelease>", lambda e, f=field: self.validate_field(e.widget, f))
            self.contract_entries[field] = entry

        self.contract_entries['date_debut'].entry.delete(0, tk.END)
        self.contract_entries['date_debut'].entry.insert(0, datetime.datetime.now().strftime("%d/%m/%Y"))
        self.contract_entries['salaire'].insert(0, "")
        self.contract_entries['prime'].insert(0, "")

        ttk.Label(contract_frame, text="Type de Salaire*", font=('Segoe UI', 10)).grid(row=len(contract_fields) + 1,
                                                                                       column=0, padx=5, pady=5,
                                                                                       sticky=tk.E)
        salary_type_frame = ttk.Frame(contract_frame)
        salary_type_frame.grid(row=len(contract_fields) + 1, column=1, sticky=tk.W)
        ttk.Radiobutton(salary_type_frame, text=" الساعة", variable=self.variables["salary_type"], value="في الساعة",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(salary_type_frame, text=" الشهر ", variable=self.variables["salary_type"], value="في الشهر ",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)

        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(fill=tk.X, pady=10)
        ttk.Button(button_frame, text="Enregistrer et Générer", command=self.save_and_generate, bootstyle=SUCCESS).pack(
            side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Réinitialiser", command=self.clear_form, bootstyle=WARNING).pack(side=tk.LEFT,
                                                                                                        padx=5)

        scrollable_frame.columnconfigure(1, weight=1)

        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind_all("<MouseWheel>", on_mousewheel)

    def calculate_fperiode(self, event=None):
        """Calcule automatiquement la date de fin (Fperiode) basée sur la durée en arabe"""
        try:
            # Vérifier si c'est un CDI
            if self.variables["contract_type"].get() != "CDI":
                return

            # Récupérer la date de début et la durée
            date_debut_str = self.get_widget_value(self.contract_entries['date_debut'])
            duree_str = self.get_widget_value(self.entries['duree'])

            if not date_debut_str or not duree_str:
                return

            # Convertir la date de début
            try:
                date_debut = datetime.datetime.strptime(date_debut_str, "%d/%m/%Y")
            except ValueError:
                return

            # Analyser la durée en arabe
            mois = 0
            annees = 0

            # Détection des mois (شهر, اشهر, أشهر)
            if any(term in duree_str for term in ["شهر", "اشهر", "أشهر"]):
                # Extraire le nombre
                numbers = re.findall(r'\d+', duree_str)
                if numbers:
                    mois = int(numbers[0])
                else:
                    # Si pas de nombre, supposer 1 mois
                    mois = 1

            # Détection des années (عام, سنة, سنوات, اعوام)
            elif any(term in duree_str for term in ["عام", "سنة", "سنوات", "اعوام"]):
                numbers = re.findall(r'\d+', duree_str)
                if numbers:
                    annees = int(numbers[0])
                else:
                    # Vérifier les termes spécifiques sans chiffres
                    if "عامين" in duree_str or "سنتين" in duree_str:
                        annees = 2
                    else:
                        annees = 1

            # Si aucune unité détectée mais il y a des chiffres, supposer des mois
            elif re.search(r'\d+', duree_str):
                numbers = re.findall(r'\d+', duree_str)
                if numbers:
                    mois = int(numbers[0])

            # Calculer la date de fin seulement si une durée a été détectée
            if mois > 0 or annees > 0:
                from dateutil.relativedelta import relativedelta
                date_fin = date_debut + relativedelta(years=annees, months=mois)

                # Mettre à jour le champ Fperiode
                self.entries['fperiode'].delete(0, tk.END)
                self.entries['fperiode'].insert(0, date_fin.strftime("%d/%m/%Y"))

                # Mettre aussi à jour le champ date_fin du contrat
                self.contract_entries['date_fin'].entry.config(state=tk.NORMAL)
                self.contract_entries['date_fin'].entry.delete(0, tk.END)
                self.contract_entries['date_fin'].entry.insert(0, date_fin.strftime("%d/%m/%Y"))

                # Re-désactiver si c'est un CDI
                if self.variables["contract_type"].get() == "CDI":
                    self.contract_entries['date_fin'].entry.config(state=tk.DISABLED)

                self.status_var.set(f"Date de fin calculée: {date_fin.strftime('%d/%m/%Y')}")

        except Exception as e:
            # CORRECTION: Utiliser logging.error au lieu de print
            logging.error(f"Erreur dans calculate_fperiode: {e}")
            self.status_var.set("Erreur dans le calcul de la durée")

    def create_search_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Recherche")

        search_frame = ttk.Frame(frame)
        search_frame.pack(fill=X, padx=10, pady=10)

        ttk.Label(search_frame, text="Matricule:", font=('Segoe UI', 10)).pack(side=LEFT)

        self.search_combo = ttk.Combobox(search_frame, textvariable=self.variables["matricule"], font=('Segoe UI', 10))
        self.search_combo.pack(side=LEFT, padx=5, expand=True, fill=X)

        ttk.Button(search_frame, text="Rechercher", command=self.search_employee, bootstyle=INFO).pack(side=LEFT,
                                                                                                       padx=5)

        info_frame = ttk.LabelFrame(frame, text="Informations Employé", bootstyle=PRIMARY)
        info_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)

        self.info_text = tk.Text(info_frame, wrap=WORD, height=12, font=('Segoe UI', 10))
        scrollbar = ttk.Scrollbar(info_frame, command=self.info_text.yview, bootstyle=PRIMARY)
        self.info_text.config(yscrollcommand=scrollbar.set)
        self.info_text.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar.pack(side=RIGHT, fill=Y)

        button_frame = ttk.Frame(frame)
        button_frame.pack(fill=X, padx=10, pady=10)

        button_frame = ttk.Frame(frame)
        button_frame.pack(fill=X, padx=10, pady=10)

        self.generate_contract_btn = ttk.Button(
            button_frame,
            text="Générer Contrat",
            command=self.generate_and_show_contract,
            bootstyle=SUCCESS,
            state=DISABLED  # Désactivé par défaut
        )
        self.generate_contract_btn.pack(side=LEFT, padx=5)

        self.view_contract_btn = ttk.Button(button_frame, text="Voir Contrat", command=self.show_last_contract,
                                            bootstyle=(PRIMARY, OUTLINE), state=DISABLED)
        self.view_contract_btn.pack(side=LEFT, padx=5)

        self.edit_btn = ttk.Button(button_frame, text="Modifier Employé",
                                   command=lambda: self.edit_employee(self.current_employee['matricule']),
                                   bootstyle=(WARNING, OUTLINE), state=DISABLED)
        self.edit_btn.pack(side=LEFT, padx=5)

        self.delete_btn = ttk.Button(button_frame, text="Supprimer Employé",
                                     command=lambda: self.delete_employee(self.current_employee['matricule']),
                                     bootstyle=(DANGER, OUTLINE), state=DISABLED)
        self.delete_btn.pack(side=LEFT, padx=5)

    def create_contract_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Contrat")

        # Configurer le widget Text avec une barre de défilement
        self.contract_text = tk.Text(frame, wrap=tk.WORD, font=('Arial', 11))
        scrollbar = ttk.Scrollbar(frame, command=self.contract_text.yview)
        self.contract_text.configure(yscrollcommand=scrollbar.set)

        # Configurer le texte de droite à gauche (RTL)
        self.contract_text.tag_configure('rtl', justify='right')

        # Disposition
        self.contract_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Configure text widget for right-to-left appearance
        self.contract_text.tag_configure("rtl", justify="right")
        self.contract_text.insert(tk.END, "", "rtl")  # Apply RTL tag to text

        button_frame = ttk.Frame(frame)
        button_frame.pack(fill=X, padx=10, pady=10)
        ttk.Button(button_frame, text="Exporter Word", command=self.export_word, bootstyle=SUCCESS).pack(side=LEFT,
                                                                                                         padx=5)
        ttk.Button(button_frame, text="Copier", command=self.copy_contract, bootstyle=INFO).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="Imprimer", command=self.print_contract, bootstyle=PRIMARY).pack(side=LEFT,
                                                                                                       padx=5)

    # Ajoutez la méthode d'impression
    def print_contract(self):
        """Imprime le contrat actuellement affiché et affiche un aperçu"""
        if not self.contract_text.get(1.0, tk.END).strip():
            Messagebox.show_warning("Aucun contrat à imprimer", "Attention")
            return

        try:
            # Create temporary Word document
            temp_dir = tempfile.gettempdir()
            temp_file = os.path.join(temp_dir, f"contrat_temp_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

            # Generate Word document
            doc = self.create_contract_doc(
                self.current_employee['matricule'],
                self.contract_text.get(1.0, tk.END).strip()
            )
            doc.save(temp_file)

            try:
                # Try printing with win32print if available
                import win32print
                import win32api

                printer_name = win32print.GetDefaultPrinter()
                win32api.ShellExecute(
                    0,
                    "print",
                    temp_file,
                    f'/d:"{printer_name}"',
                    temp_dir,
                    0
                )
                self.status_var.set(f"Contrat envoyé à l'imprimante {printer_name}")

            except ImportError:
                # Fallback for systems without win32print
                if os.name == 'posix':
                    subprocess.run(['lpr', temp_file])
                    self.status_var.set("Contrat envoyé à l'imprimante par défaut")
                else:
                    os.startfile(temp_file, "print")
                    self.status_var.set("Ouverture du contrat pour impression")

            # Open the document for preview
            try:
                if os.name == 'nt':
                    os.startfile(temp_file)
                else:
                    subprocess.run(['xdg-open', temp_file])
                Messagebox.show_info(
                    f"Le contrat a été envoyé à l'imprimante et ouvert pour aperçu.\n"
                    f"Fichier temporaire: {temp_file}",
                    "Impression et Aperçu historically"
                )

            except Exception as preview_error:
                Messagebox.show_warning(
                    f"Contrat imprimé, mais erreur lors de l'ouverture de l'aperçu:\n{str(preview_error)}\n"
                    f"Fichier temporaire: {temp_file}",
                    "Avertissement"
                )

        except Exception as e:
            print(
                f"Erreur lors de la création/impression du contrat:\n{str(e)}",
                "Erreur d'impression"
            )

    def create_list_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Liste Employés")

        # Frame pour les statistiques et actions
        summary_frame = ttk.Frame(frame, bootstyle=INFO)
        summary_frame.pack(fill=X, padx=10, pady=5)

        # Statistiques principales
        stats_frame = ttk.Frame(summary_frame)
        stats_frame.pack(side=LEFT, fill=X, expand=True)

        # Labels pour les statistiques
        self.total_label = ttk.Label(stats_frame, text="Total: 0", font=('Segoe UI', 10))
        self.total_label.pack(side=LEFT, padx=10)

        self.cdd_label = ttk.Label(stats_frame, text="CDD: 0", font=('Segoe UI', 10))
        self.cdd_label.pack(side=LEFT, padx=10)

        self.cdi_label = ttk.Label(stats_frame, text="CDI: 0", font=('Segoe UI', 10))
        self.cdi_label.pack(side=LEFT, padx=10)

        self.salary_label = ttk.Label(stats_frame, text="Salaire Moyen: 0.00 TND", font=('Segoe UI', 10))
        self.salary_label.pack(side=LEFT, padx=10)

        self.prime_label = ttk.Label(stats_frame, text="Prime Moyen: 0.00 TND", font=('Segoe UI', 10))
        self.prime_label.pack(side=LEFT, padx=10)

        # Bouton d'actualisation
        refresh_btn = ttk.Button(
            summary_frame,
            text="🔄 Actualiser",
            command=self.load_employee_table,
            bootstyle=(INFO, OUTLINE)
        )
        refresh_btn.pack(side=RIGHT, padx=5)

        filter_frame = ttk.Frame(frame)
        filter_frame.pack(fill=X, padx=10, pady=5)
        ttk.Label(filter_frame, text="Filtrer par:", font=('Segoe UI', 10)).pack(side=LEFT)
        self.filter_var = tk.StringVar()
        self.filter_combo = ttk.Combobox(filter_frame, textvariable=self.filter_var,
                                         values=["Nom", "Matricule", "Type Contrat"],
                                         font=('Segoe UI', 10))
        self.filter_combo.pack(side=LEFT, padx=5)
        self.filter_entry = ttk.Entry(filter_frame, font=('Segoe UI', 10))
        self.filter_entry.pack(side=LEFT, padx=5, expand=True, fill=X)
        ttk.Button(filter_frame, text="Filtrer", command=self.apply_filter, bootstyle=INFO).pack(side=LEFT, padx=5)
        ttk.Button(filter_frame, text="Réinitialiser", command=self.reset_filter, bootstyle=WARNING).pack(side=LEFT,
                                                                                                          padx=5)
        ttk.Button(filter_frame, text="Exporter Tous", command=self.export_all_contracts, bootstyle=SUCCESS).pack(
            side=LEFT, padx=5)
        ttk.Button(filter_frame, text="Annuler", command=self.undo_action, bootstyle=(WARNING, OUTLINE)).pack(side=LEFT,
                                                                                                              padx=5)

        self.employee_table = Tableview(
            frame, coldata=self.column_definitions, rowdata=[], paginated=True, searchable=True,
            bootstyle=PRIMARY, autoalign=True, stripecolor=('lightblue', None), pagesize=20
        )
        self.employee_table.pack(fill=BOTH, expand=True, padx=10, pady=10)

        self.employee_table.view.bind("<Double-1>", self.edit_cell)
        self.context_menu = tk.Menu(self.root, tearoff=0, font=('Segoe UI', 10))
        self.context_menu.add_command(label="Modifier", command=self.context_menu_modify)
        self.context_menu.add_command(label="Supprimer", command=self.context_menu_delete)
        self.context_menu.add_command(label="Voir Contrat", command=self.context_menu_view_contract)
        self.context_menu.add_command(label="Exporter Contrat", command=self.context_menu_export_contract)
        self.employee_table.view.bind("<Button-3>", self.show_context_menu)

    def validate_field(self, widget, field):
        value = widget.get()
        validators = {
            "matricule": lambda v: bool(re.match(r'^\w+$', v)) if v else False,
            "cin": lambda v: bool(re.match(r'^\d{8}$', v)) if v else True,
            "date_naissance": lambda v: bool(re.match(r'^\d{4}/\d{2}/\d{2}$', v)) if v else True,
            "date_cin": lambda v: bool(re.match(r'^\d{4}/\d{2}/\d{2}$', v)) if v else True,
            "email": lambda v: bool(re.match(r'^[^@]+@[^@]+\.[^@]+$', v)) if v else True,
            "telephone": lambda v: bool(re.match(r'^\+?\d{10,12}$', v)) if v else True,
            "salaire": lambda v: bool(re.match(r'^\d+(\.\d{1,3})?$', v)) and float(v) > 0 if v else False,
            "prime": lambda v: bool(re.match(r'^\d+(\.\d{1,3})?$', v)) and float(v) >= 0 if v else False,
            "date_debut": lambda v: bool(re.match(r'^\d{4}/\d{2}/\d{2}$', v)) if v else False,
            "date_fin": lambda v: bool(re.match(r'^\d{4}/\d{2}/\d{2}$', v)) if v else True
        }
        widget.configure(bootstyle="danger" if not validators.get(field, lambda x: True)(value) else "primary")

    def toggle_date_fin(self):
        contract_type = self.variables["contract_type"].get()
        state = DISABLED if contract_type == "CDI" else NORMAL
        self.contract_entries['date_fin'].entry.config(state=state)

        # Si c'est un CDI, calculer automatiquement Fperiode si une durée existe
        if contract_type == "CDI":
            self.calculate_fperiode()

    def load_data(self):
        self.load_matricules()
        self.load_employee_table()
        # Vérifier les contrats proches d'expiration après un court délai

    def load_matricules(self):
        with self.conn.cursor() as cursor:
            cursor.execute("SELECT matricule FROM employees ORDER BY matricule")
            self.search_combo['values'] = [row[0] for row in cursor.fetchall()]

    def apply_filter(self):
        filter_type = self.filter_var.get()
        filter_value = self.filter_entry.get().strip().lower()
        if not filter_type or not filter_value:
            self.load_employee_table()
            return

        query = '''
                SELECT e.matricule, \
                       e.nom, \
                       e.prenom, \
                       e.genre, \
                       e.date_naissance, \
                       e.lieu_naissance,
                       e.adresse, \
                       e.ville, \
                       e.cin, \
                       e.date_cin, \
                       e.lieu_cin, \
                       e.poste, \
                       e.email, \
                       e.telephone,
                       c.type_contrat, \
                       c.date_debut, \
                       c.date_fin, \
                       c.salaire_base, \
                       c.prime, \
                       c.salary_type
                FROM employees e
                         LEFT JOIN contrats c ON e.matricule = c.matricule
                    AND c.id = (SELECT MAX(id) FROM contrats WHERE matricule = e.matricule)
                WHERE {} \
                '''
        params = (f"%{filter_value}%",) if filter_type != "Type Contrat" else (filter_value.upper(),)
        condition = {
            "Nom": "LOWER(e.nom) LIKE %s",
            "Matricule": "e.matricule LIKE %s",
            "Type Contrat": "c.type_contrat = %s"
        }.get(filter_type, "1=1")

        with self.conn.cursor() as cursor:
            cursor.execute(query.format(condition), params)
            self.update_table_data(cursor.fetchall())
            self.status_var.set(f"{len(self.employee_table.get_rows())} employés trouvés")

    def reset_filter(self):
        self.filter_var.set("")
        self.filter_entry.delete(0, tk.END)
        self.load_employee_table()

    def show_context_menu(self, event):
        row_id = self.employee_table.view.identify_row(event.y)
        if row_id:
            self.employee_table.view.selection_set(row_id)
            self.selected_matricule = self.employee_table.get_row(row_id).values[0]
            self.context_menu.post(event.x_root, event.y_root)

    def context_menu_modify(self):
        self.edit_employee(self.selected_matricule)

    def context_menu_delete(self):
        self.delete_employee(self.selected_matricule)

    def context_menu_view_contract(self):
        self.current_employee = {'matricule': self.selected_matricule}
        self.show_last_contract()

    def context_menu_export_contract(self):
        self.current_employee = {'matricule': self.selected_matricule}
        self.show_last_contract()
        self.export_word()

    def clear_form(self):
        for entry in self.entries.values():
            if isinstance(entry, ttk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):  # Handle DateEntry differently
                entry.entry.delete(0, tk.END)

        for entry in self.contract_entries.values():
            if isinstance(entry, ttk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):  # Handle DateEntry differently
                entry.entry.delete(0, tk.END)
        self.entries['ville'].insert(0, "المحرس")
        self.entries['lieu_cin'].insert(0, "تونس")
        self.contract_entries['date_debut'].entry.insert(0, datetime.datetime.now().strftime("%d/%m/%Y"))
        self.contract_entries['salaire'].insert(0, "2500.00")
        self.contract_entries['prime'].insert(0, "500.00")
        self.variables["genre"].set("féminin")
        self.variables["contract_type"].set("CDD")
        self.variables["salary_type"].set("في الساعة")
        self.status_var.set("Formulaire réinitialisé")

    def view_contract_from_table(self, matricule):
        self.current_employee = {'matricule': matricule}
        self.show_last_contract()

    def export_word(self):
        """Exporte le contrat au format Word en ajoutant le matricule dans le document."""
        try:
            # Vérification de l'employé sélectionné (votre code original)
            if not getattr(self, 'current_employee', None):
                print("Aucun employé sélectionné.", "Erreur", parent=self.root)
                return

            emp = self.current_employee
            contract_type = emp.get('type_contrat', 'CDD').upper()

            if contract_type not in ['CDD', 'CDI']:
                print("Type de contrat invalide. Doit être 'CDD' ou 'CDI'.", "Erreur", parent=self.root)
                return

            # Chemins des templates (votre code original inchangé)
            template_paths = {
                'CDD': r"D:\CDD CONTRAT.docx",
                'CDI': r"D:\CDI CONTRAT.docx"
            }

            template_path = template_paths.get(contract_type)
            if not template_path or not os.path.exists(template_path):
                print(f"Le modèle {contract_type} est introuvable : {template_path}",
                                      "Fichier manquant", parent=self.root)
                return

            # Nom de fichier (votre code original inchangé)
            filename = f"Contrat_{contract_type}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

            # Votre code original pour l'export
            save_path = filedialog.asksaveasfilename(
                title=f"Enregistrer le contrat {contract_type}",
                defaultextension=".docx",
                initialfile=filename,
                filetypes=[("Document Word", "*.docx"), ("Tous les fichiers", "*.*")]
            )
            if not save_path:
                return

            # Copie du template (votre code original)
            shutil.copy2(template_path, save_path)

            # Chargement du document
            doc = Document(save_path)

            # AJOUT: Insertion du matricule en bas de l'en-tête
            matricule = emp.get('matricule', 'N/C')
            header_section = doc.sections[0]
            header = header_section.header

            # Création d'un paragraphe pour le matricule
            matricule_para = header.add_paragraph()
            matricule_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            matricule_run = matricule_para.add_run(f"{matricule}")
            matricule_run.font.size = Pt(9)
            matricule_run.bold = True

            # Votre code original pour le corps du document
            style_name = 'ArabicStyle'
            if style_name not in doc.styles:
                arabic_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                arabic_style.font.name = 'Arial'
                arabic_style.font.size = Pt(11)
                arabic_style.font.rtl = True
            else:
                arabic_style = doc.styles[style_name]

            contrat_text = self.contract_text.get("1.0", tk.END).strip()
            for line in contrat_text.splitlines():
                if line.strip():
                    p = doc.add_paragraph(line.strip(), style=style_name)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    p.paragraph_format.space_after = Pt(6)

            # Sauvegarde finale
            doc.save(save_path)

            # Message de confirmation (votre code original)
            Messagebox.show_info(f"Contrat {contract_type} généré avec succès.", "Export réussi", parent=self.root)
            self.status_var.set(f"Contrat {contract_type} exporté")

            # Ouverture du document (votre code original)
            if Messagebox.yesno(f"Contrat {contract_type} généré avec succès.\nSouhaitez-vous l'ouvrir ?",
                                "Export réussi", parent=self.root):
                os.startfile(save_path)

        except PermissionError:
            print("Impossible d'accéder au fichier. Veuillez fermer Word et réessayer.",
                                  "Erreur d'accès", parent=self.root)
        except Exception as e:
            print(f"Erreur lors de l'export :\n{str(e)}", "Erreur d'export", parent=self.root)

################################################################################################

    def create_progress_dialog(self, title, max_value):
        """Create a modal progress dialog with a progress bar."""
        self.progress_dialog = ttk.Toplevel(self.root)
        self.progress_dialog.title(title)
        self.progress_dialog.geometry("400x150")
        self.progress_dialog.transient(self.root)
        self.progress_dialog.grab_set()

        # Label for progress percentage
        self.progress_label = ttk.Label(self.progress_dialog, text="0% (0/0 contrats exportés)")
        self.progress_label.pack(pady=20)

        # Progress bar
        self.progress_bar = ttk.Progressbar(
            self.progress_dialog,
            maximum=max_value,
            value=0,
            length=350,
            mode='determinate'
        )
        self.progress_bar.pack(pady=10)

        # Center dialog
        self.progress_dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - self.progress_dialog.winfo_width()) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - self.progress_dialog.winfo_height()) // 2
        self.progress_dialog.geometry(f"+{x}+{y}")

    def update_progress(self, current, total):
        """Update the progress bar and label."""
        if self.progress_bar and self.progress_dialog:
            percentage = (current / total) * 100
            self.progress_bar['value'] = current
            self.progress_label.config(text=f"{percentage:.1f}% ({current}/{total} contrats exportés)")
            self.root.update()

    def close_progress_dialog(self):
        """Close the progress dialog."""
        if self.progress_dialog:
            self.progress_dialog.grab_release()
            self.progress_dialog.destroy()
            self.progress_dialog = None
            self.progress_bar = None
            self.progress_label = None

    def export_all_contracts(self):
        try:
            with self.conn.cursor() as cursor:
                # Count total contracts to set progress bar maximum
                cursor.execute("""
                               SELECT COUNT(DISTINCT matricule)
                               FROM contrats
                               """)
                total_contracts = cursor.fetchone()[0]

                if total_contracts == 0:
                    Messagebox.show_info("Aucun contrat à exporter", "Information", parent=self.root)
                    return

                # Create progress dialog
                self.create_progress_dialog("Exportation des contrats", total_contracts)

                # Fetch latest contracts
                cursor.execute("""
                               SELECT c.matricule, c.texte_contrat
                               FROM contrats c
                                        INNER JOIN (SELECT matricule, MAX(date_creation) as max_date
                                                    FROM contrats
                                                    GROUP BY matricule) latest
                                                   ON c.matricule = latest.matricule AND c.date_creation = latest.max_date
                               """)
                contracts = cursor.fetchall()

                folder = filedialog.askdirectory(title="Choisir le dossier de destination")
                if not folder:
                    self.close_progress_dialog()
                    return

                zip_path = os.path.join(folder,
                                        f"contrats_export_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip")
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for i, contract in enumerate(contracts, 1):
                        matricule, texte_contrat = contract
                        doc = self.create_contract_doc(matricule, texte_contrat)
                        temp_file = os.path.join(folder, f"contrat_{matricule}.docx")
                        doc.save(temp_file)
                        zipf.write(temp_file, os.path.basename(temp_file))
                        os.remove(temp_file)
                        # Update progress
                        self.update_progress(i, total_contracts)

                self.close_progress_dialog()
                Messagebox.show_info(f"Contrats exportés avec succès dans {zip_path}", "Succès", parent=self.root)
        except Exception as e:
            self.close_progress_dialog()
            print(f"Erreur lors de l'exportation: {str(e)}", "Erreur", parent=self.root)

################################################################################################

    def copy_contract(self):
        self.root.clipboard_clear()
        self.root.clipboard_append(self.contract_text.get(1.0, tk.END))
        self.status_var.set("Texte du contrat copié")

    def edit_cell(self, event):
        # Récupérer la ligne et la colonne cliquées
        row_id = self.employee_table.view.identify_row(event.y)
        column = self.employee_table.view.identify_column(event.x)

        if not row_id or not column:
            return

        # Convertir l'ID de ligne en index numérique
        try:
            row_index = int(row_id.replace('I', '')) - 1  # Convertir 'I001' en 0, etc.
        except ValueError:
            return

        # Vérifier si l'index est valide
        if row_index < 0 or row_index >= len(self.employee_table.tablerows):
            return

        col_idx = int(column.replace("#", "")) - 1
        col_name = self.column_definitions[col_idx]["text"]

        # Empêcher l'édition de certaines colonnes
        if col_name in ["Matricule", "Actions"]:
            return

        # Récupérer les données de la ligne
        row_data = self.employee_table.tablerows[row_index].values
        matricule = row_data[0]
        current_value = row_data[col_idx]

        # Créer le champ d'édition
        entry = ttk.Entry(self.employee_table.view, bootstyle="primary", font=('Segoe UI', 10))
        entry.insert(0, current_value)

        # Positionner le champ d'édition
        cell_bbox = self.employee_table.view.bbox(row_id, column)
        if cell_bbox:
            entry.place(x=cell_bbox[0], y=cell_bbox[1], width=cell_bbox[2], height=cell_bbox[3])

        def validate_input(value):
            validators = {
                "Date Naissance": r'^\d{2}/\d{2}/\d{4}$',
                "Date CIN": r'^\d{2}/\d{2}/\d{4}$',
                "Date Début": r'^\d{2}/\d{2}/\d{4}$',
                "Date Fin": r'^\d{2}/\d{2}/\d{4}$',
                "Email": r'^[^@]+@[^@]+\.[^@]+$',
                "Genre": r'^(féminin|masculin)$',
                "Type Contrat": r'^(CDD|CDI)$',
                "Salaire Base": r'^\d+(\.\d{1,2})?$',
                "Prime": r'^\d+(\.\d{1,2})?$',
                "Type Salaire": r'^(hourly|monthly)$',
                "Nb Échéances": r'^\d+$'
            }
            pattern = validators.get(col_name, r'.*')
            if not re.match(pattern, value):
                return False
            if col_name in ["Salaire Base", "Prime"] and value:
                try:
                    return float(value) > 0
                except ValueError:
                    return False
            return True

        def save_edit(event=None):
            new_value = entry.get().strip()
            if not validate_input(new_value):
                Messagebox.show_warning(f"Valeur invalide pour {col_name}", "Erreur")
                entry.destroy()
                return

            try:
                with self.conn.cursor() as cursor:
                    sql_field = {
                        "Nom": "nom", "Prénom": "prenom", "Genre": "genre",
                        "Date Naissance": "date_naissance", "Lieu Naissance": "lieu_naissance",
                        "Adresse": "adresse", "Ville": "ville", "CIN": "cin",
                        "Date CIN": "date_cin", "Lieu CIN": "lieu_cin", "Poste": "poste",
                        "Email": "email", "Téléphone": "telephone", "Type Contrat": "type_contrat",
                        "Date Début": "date_debut", "Date Fin": "date_fin",
                        "Salaire Base": "salaire_base", "Prime": "prime",
                        "Type Salaire": "salary_type", "Atelier": "atelier",
                        "Nb Échéances": "nbre_eche"
                    }.get(col_name)

                    if sql_field:
                        if col_name in ["Type Contrat", "Date Début", "Date Fin", "Salaire Base", "Prime",
                                        "Type Salaire"]:
                            # Mise à jour dans la table contrats
                            cursor.execute("SELECT id FROM contrats WHERE matricule = %s ORDER BY id DESC LIMIT 1",
                                           (matricule,))
                            contract_id = cursor.fetchone()
                            if contract_id:
                                cursor.execute(f"UPDATE contrats SET {sql_field} = %s WHERE id = %s",
                                               (float(new_value) if col_name in ["Salaire Base",
                                                                                 "Prime"] else new_value,
                                                contract_id[0]))
                                self.undo_stack.append(
                                    ("contract_update", matricule, contract_id[0], sql_field, current_value))
                        else:
                            # Mise à jour dans la table employees
                            cursor.execute(f"UPDATE employees SET {sql_field} = %s WHERE matricule = %s",
                                           (new_value, matricule))
                            self.undo_stack.append(("employee_update", matricule, sql_field, current_value))

                        self.conn.commit()
                        self.load_employee_table()
                        self.status_var.set(f"Champ {col_name} mis à jour pour {matricule}")

                        if self.current_employee and self.current_employee['matricule'] == matricule:
                            self.search_employee()

            except Exception as e:
                Messagebox.showerror("Erreur", f"Erreur de mise à jour: {str(e)}")
            finally:
                entry.destroy()

        entry.bind("<Return>", save_edit)
        entry.bind("<FocusOut>", save_edit)
        entry.focus_set()


    def validate_field_value(self, field, value):
        validators = {
            "matricule": lambda v: bool(re.match(r'^\w+$', v)) if v else False,
            "cin": lambda v: bool(re.match(r'^\d{8}$', v)) if v else True,
            "date_naissance": lambda v: self.is_valid_date(v) if v else True,
            "date_cin": lambda v: self.is_valid_date(v) if v else True,
            "email": lambda v: bool(re.match(r'^[^@]+@[^@]+\.[^@]+$', v)) if v else True,
            "telephone": lambda v: bool(re.match(r'^\+?\d{10,12}$', v)) if v else True,
            "salaire": lambda v: bool(re.match(r'^\d+(\.\d{1,2})?$', v)) and float(v) > 0 if v else False,
            "prime": lambda v: bool(re.match(r'^\d+(\.\d{1,2})?$', v)) and float(v) >= 0 if v else False,
            "date_debut": lambda v: self.is_valid_date(v) if v else False,
            "date_fin": lambda v: self.is_valid_date(v) if v else True
        }
        return validators.get(field, lambda x: True)(value)

    def is_valid_date(self, date_str):
        try:
            if not date_str or not re.match(r'^\d{4}/\d{2}/\d{2}$', date_str.strip()):
                return False
            datetime.datetime.strptime(date_str.strip(), "%d/%m/%Y")
            return True
        except ValueError:
            return False



    def delete_employee(self, matricule):
        # Vérification que le matricule est valide (uniquement des chiffres)
        if not matricule or not matricule.isdigit():
            print("Matricule invalide", "Erreur")
            return

        # Confirmation de la suppression
        if not Messagebox.yesno(f"Confirmer la suppression de l'employé avec matricule {matricule} ?",
                                "Confirmation de suppression"):
            return

        try:
            with self.conn.cursor() as cursor:
                # Vérifier d'abord si l'employé existe
                cursor.execute("SELECT 1 FROM employees WHERE matricule = %s", (matricule,))
                if not cursor.fetchone():
                    print(f"Aucun employé trouvé avec le matricule {matricule}", "Erreur")
                    return

                # Sauvegarde des données pour undo (optionnel)
                cursor.execute("SELECT * FROM employees WHERE matricule = %s", (matricule,))
                employee_data = cursor.fetchone()

                # Suppression des contrats associés
                cursor.execute("DELETE FROM contrats WHERE matricule = %s", (matricule,))

                # Suppression de l'employé
                cursor.execute("DELETE FROM employees WHERE matricule = %s", (matricule,))

                self.conn.commit()

                # Mise à jour de l'interface
                self.load_data()
                self.clear_search()

                # Message de confirmation
                self.status_var.set(f"Employé {matricule} supprimé avec succès")
                Messagebox.show_info(f"L'employé avec matricule {matricule} a été supprimé", "Succès")

        except pymysql.MySQLError as e:
            self.conn.rollback()
            print(f"Erreur lors de la suppression : {str(e)}", "Erreur")
            self.status_var.set("Erreur de suppression")

    def undo_action(self):
        if not self.undo_stack:
            Messagebox.show_info("Aucune action à annuler", "Information")
            return

        action_type, matricule, employee_data, contract_data = self.undo_stack.pop()
        try:
            with self.conn.cursor() as cursor:
                if action_type == "employee_delete":
                    cursor.execute('''
                                   INSERT INTO employees (matricule, nom, prenom, genre, date_naissance, lieu_naissance,
                                                          adresse, ville, cin, date_cin, lieu_cin, poste, email,
                                                          telephone)
                                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                   ''', (
                                       employee_data['matricule'], employee_data['nom'], employee_data['prenom'],
                                       employee_data['genre'],
                                       employee_data['date_naissance'], employee_data['lieu_naissance'],
                                       employee_data['adresse'],
                                       employee_data['ville'], employee_data['cin'], employee_data['date_cin'],
                                       employee_data['lieu_cin'],
                                       employee_data['poste'], employee_data['email'], employee_data['telephone']
                                   ))
                    if contract_data:
                        cursor.execute('''
                                       INSERT INTO contrats ( matricule, type_contrat, date_creation, date_debut,
                                                             date_fin, salaire_base, prime, salary_type, texte_contrat)
                                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                       ''', (
                                           contract_data['id'], matricule, contract_data['type_contrat'],
                                           contract_data['date_creation'],
                                           contract_data['date_debut'], contract_data['date_fin'],
                                           contract_data['salaire_base'],
                                           contract_data['prime'], contract_data['salary_type'],
                                           contract_data['texte_contrat']
                                       ))
                    self.status_var.set(f"Suppression de {matricule} annulée")
                elif action_type == "employee_update":
                    cursor.execute(f"UPDATE employees SET {employee_data} = %s WHERE matricule = %s",
                                   (contract_data, matricule))
                    self.status_var.set(f"Mise à jour de {employee_data} pour {matricule} annulée")
                elif action_type == "contract_update":
                    cursor.execute(f"UPDATE contrats SET {employee_data} = %s WHERE id = %s",
                                   (contract_data, matricule))
                    self.status_var.set(f"Mise à jour du contrat pour {matricule} annulée")

                self.conn.commit()
                self.load_data()
                if self.current_employee and self.current_employee['matricule'] == matricule:
                    self.search_employee()
        except pymysql.MySQLError as e:
            print(f"Erreur lors de l'annulation: {str(e)}", "Erreur")

    def show_help(self):
        Messagebox.show_info(
            """Application de Gestion des Contrats
    Version: 1.0
    Fonctionnalités:
    - Ajouter, modifier, supprimer des employés
    - Générer des contrats CDD/CDI en arabe
    - Exporter les contrats en Word ou ZIP
    - Rechercher et filtrer les employés
    - Modifier les données directement dans le tableau
    - Annuler la dernière action (suppression ou modification)

    Pour plus d'aide, contactez le support technique.""",
            "Aide"
        )

    def show_about(self):
        Messagebox.show_info(
            """Application de Gestion des Contrats
    Développée par: Omar Badrani
    Version: 1.0
    © Imbert Mnif. Tous droits réservés.""",
            "À propos"
        )

##########################################################################################################


#####################################################################################################
    def on_alert_double_click(self, event):
        """Gérer le double-clic sur une ligne du tableau des alertes pour éditer un employé."""
        try:
            # Obtenir la sélection actuelle
            selection = self.alert_table.view.selection()
            if not selection:
                Messagebox.show_warning("Aucune ligne sélectionnée", "Attention", parent=self.root)
                return

            # Prendre la première ligne sélectionnée
            selected_item = selection[0]

            # Obtenir les valeurs de la ligne sélectionnée
            item_values = self.alert_table.view.item(selected_item, 'values')
            if not item_values:
                print("Impossible de lire les données de la ligne", "Erreur", parent=self.root)
                return

            # Le matricule est la première colonne
            matricule = item_values[0]

            # Vérifier la validité du matricule
            if not matricule or not str(matricule).strip():
                print("Matricule invalide", "Erreur", parent=self.root)
                return

            # Ouvrir la fenêtre d'édition
            self.edit_employee(matricule)

        except Exception as e:
            print(f"Erreur inattendue : {str(e)}", "Erreur", parent=self.root)

    def create_alerts_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Alertes Contrats")

        alert_frame = ttk.LabelFrame(frame, text="Contrats Expirant dans 30 Jours ou Moins", bootstyle="primary")
        alert_frame.pack(fill=BOTH, expand=True, padx=10, pady=10)

        self.alert_table = Tableview(
            alert_frame,
            coldata=self.alert_column_definitions,
            rowdata=[],
            paginated=True,
            searchable=True,
            bootstyle="primary",
            autoalign=True,
            stripecolor=('lightblue', None),
            pagesize=20
        )
        self.alert_table.pack(fill=BOTH, expand=True, padx=10, pady=10)

        # Ajouter le binding pour le double-clic (version simplifiée)
        self.alert_table.view.bind("<Double-1>", self.on_alert_double_click)

        button_frame = ttk.Frame(alert_frame)
        button_frame.pack(fill=X, pady=10)
        ttk.Button(button_frame, text="Actualiser", command=self.load_alert_table, bootstyle="info").pack(side=LEFT,
                                                                                                          padx=5)
        ttk.Button(button_frame, text="Effacer Alertes", command=self.clear_alerts, bootstyle="primary").pack(side=LEFT,
                                                                                                              padx=5)

    def edit_selected_alert_employee(self):
        """Méthode de secours pour éditer l'employé sélectionné dans les alertes"""
        try:
            # Obtenir la sélection actuelle
            selection = self.alert_table.view.selection()
            if not selection:
                Messagebox.show_warning("Aucun employé sélectionné", "Attention", parent=self.root)
                return

            # Prendre la première ligne sélectionnée
            selected_item = selection[0]

            # Obtenir les valeurs de la ligne sélectionnée
            item_values = self.alert_table.view.item(selected_item, 'values')
            if not item_values:
                print("Impossible de lire les données de la ligne", "Erreur", parent=self.root)
                return

            # Le matricule est la première colonne
            matricule = item_values[0]

            # Vérifier la validité du matricule
            if not matricule or not str(matricule).strip():
                print("Matricule invalide", "Erreur", parent=self.root)
                return

            # Ouvrir la fenêtre d'édition
            self.edit_employee(matricule)

        except Exception as e:
            print(f"Erreur inattendue : {str(e)}", "Erreur", parent=self.root)


    def clear_alerts(self):
        self.alerted_contracts.clear()
        self.load_alert_table()
        self.status_var.set("Alertes effacées")

    def check_expiring_contracts(self, force_notification=False):
        """Vérifie les contrats sur le point d'expirer"""
        try:
            today = datetime.datetime.now().date().strftime('%Y-%m-%d')

            with self.conn.cursor() as cursor:
                cursor.execute('''
                               SELECT e.matricule,
                                      e.nom,
                                      e.prenom,
                                      e.date_fin,
                                      DATEDIFF(STR_TO_DATE(e.date_fin, '%%Y-%%m-%%d'),
                                               STR_TO_DATE(%s, '%%Y-%%m-%%d')) AS jours_restants
                               FROM employees e
                               WHERE e.type_contrat = 'CDD'
                                 AND e.date_fin IS NOT NULL
                                 AND e.date_fin != ''
                      AND DATEDIFF(STR_TO_DATE(e.date_fin, '%%Y-%%m-%%d'), 
                           STR_TO_DATE(%s
                                   , '%%Y-%%m-%%d')) BETWEEN 0
                                 AND 30
                               ORDER BY jours_restants
                               ''', (today, today))

                expiring_contracts = cursor.fetchall()

            new_alerts = []
            for contract in expiring_contracts:
                matricule, nom, prenom, date_fin, jours_restants = contract
                contract_key = f"{matricule}_{date_fin}"

                # Vérifie si le contrat a changé ou si l'alerte est forcée
                if contract_key not in self.alerted_contracts or force_notification:
                    new_alerts.append(contract)
                    self.alerted_contracts[contract_key] = {
                        'date_fin': date_fin,
                        'last_alerted': datetime.datetime.now()
                    }
                else:
                    # Vérifie si on doit répéter l'alerte (toutes les 24h)
                    last_alert = self.alerted_contracts[contract_key]['last_alerted']
                    if (datetime.datetime.now() - last_alert).total_seconds() >= 24 * 3600:
                        new_alerts.append(contract)
                        self.alerted_contracts[contract_key]['last_alerted'] = datetime.datetime.now()

            if new_alerts:
                self.show_contract_alerts(new_alerts)
                self.load_alert_table()

        except Exception as e:
            logging.error(f"Erreur lors de la vérification des contrats: {str(e)}")
            self.status_var.set("Erreur vérification alertes")

    def show_contract_alerts(self, contracts):
        """Affiche les alertes de contrat"""
        message = "⚠️ ALERTE : Contrats CDD expirant bientôt ⚠️\n\n"
        for contract in contracts:
            matricule, nom, prenom, date_fin, jours_restants = contract
            message += f"• {nom} {prenom} (Matricule: {matricule}) - "
            message += f"Expire le {date_fin} (dans {jours_restants} jours)\n"

        self.play_alert_sound()
        Messagebox.show_warning(message, "Alerte Contrats", parent=self.root)
        self.stop_alert_sound()
        self.status_var.set(f"⚠ {len(contracts)} nouveaux contrats expirent bientôt")

    def play_alert_sound(self):
        if self.sound_enabled and hasattr(self, 'sound_file') and self.sound_file and os.path.exists(self.sound_file):
            try:
                if os.name == 'nt':
                    def play_loop():
                        while not hasattr(self, 'alert_stopped') or not self.alert_stopped:
                            winsound.PlaySound(self.sound_file, winsound.SND_FILENAME | winsound.SND_ASYNC)
                            time.sleep(2)
                else:
                    def play_loop():
                        while not hasattr(self, 'alert_stopped') or not self.alert_stopped:
                            subprocess.run(['aplay', self.sound_file], check=False)
                            time.sleep(2)

                self.alert_thread = threading.Thread(target=play_loop)
                self.alert_thread.daemon = True
                self.alert_thread.start()
            except Exception as e:
                logging.error(f"Erreur de lecture du son d'alerte: {e}")

    def stop_alert_sound(self):
        """Arrête le son d'alerte"""
        try:
            self.alert_stopped = True
            winsound.PlaySound(None, 0)
            if hasattr(self, 'alert_thread'):
                self.alert_thread.join(timeout=0.1)
        except Exception as e:
            print(f"Erreur lors de l'arrêt du son: {e}")

    def stop_alert_timer(self):
        if self.alert_timer:
            self.root.after_cancel(self.alert_timer)

    def update_summary(self):
        try:
            with self.conn.cursor() as cursor:
                # Compter le nombre total d'employés
                cursor.execute("SELECT COUNT(*) FROM employees")
                total_employees = cursor.fetchone()[0] or 0

                # Compter les CDD et CDI
                cursor.execute("""
                               SELECT SUM(CASE WHEN type_contrat = 'CDD' THEN 1 ELSE 0 END) as cdd_count,
                                      SUM(CASE WHEN type_contrat = 'CDI' THEN 1 ELSE 0 END) as cdi_count,
                                      AVG(salaire_base)as avg_salary,
                                      AVG(prime) as avg_prime
                               FROM employees
                               """)
                stats = cursor.fetchone()

                cdd_count = stats[0] or 0
                cdi_count = stats[1] or 0
                avg_salary = stats[2] or 0
                avg_prime=stats[3] or 0
                # Mettre à jour l'interface
                self.total_label.config(text=f"Total: {total_employees}")
                self.cdd_label.config(text=f"CDD: {cdd_count}")
                self.cdi_label.config(text=f"CDI: {cdi_count}")
                self.salary_label.config(text=f"Salaire Moyen: {avg_salary:.2f} TND")

                self.prime_label.config(text=f"Prime Moyen: {avg_prime:.2f} TND")

        except pymysql.Error as e:
            print(f"Erreur lors de la mise à jour du résumé: {e}")
            self.status_var.set("Erreur statistiques")

    def edit_employee(self, matricule):
        """Open a window to edit an employee's details with a scrollable interface."""
        top = ttk.Toplevel(self.root)
        top.title(f"Modifier Employé {matricule}")
        top.geometry("900x700")  # Augmenter la hauteur pour les nouveaux champs

        try:
            with self.conn.cursor() as cursor:
                cursor.execute('''
                    SELECT matricule, nom, prenom, genre, date_naissance, lieu_naissance,
                           adresse, ville, cin, date_cin, lieu_cin, poste, email, telephone,
                           type_contrat, date_debut, date_fin, salaire_base, prime, salary_type,
                           dcon, duree, atelier, nbre_eche, fperiode
                    FROM employees
                    WHERE matricule = %s
                ''', (matricule,))
                employee = cursor.fetchone()

                cursor.execute('''
                    SELECT type_contrat, date_creation, texte_contrat
                    FROM contrats
                    WHERE matricule = %s
                    ORDER BY date_creation DESC LIMIT 1
                ''', (matricule,))
                contract = cursor.fetchone()

            if not employee:
                Messagebox.show_error("Employé non trouvé", "Erreur", parent=top)
                top.destroy()
                return
        except Exception as e:
            Messagebox.show_error(f"Erreur de base de données: {str(e)}", "Erreur", parent=top)
            top.destroy()
            return

        # Créer l'interface avec les nouveaux champs
        canvas = tk.Canvas(top, highlightthickness=0)
        scrollbar = ttk.Scrollbar(top, orient="vertical", command=canvas.yview, bootstyle=PRIMARY)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        notebook = ttk.Notebook(scrollable_frame, bootstyle=PRIMARY)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        employee_frame = ttk.Frame(notebook)
        contract_frame = ttk.Frame(notebook)
        notebook.add(employee_frame, text="Détails Employé")
        notebook.add(contract_frame, text="Détails Contrat")

        genre_var = tk.StringVar(value=employee[3] if employee[3] else "السيدة")
        contract_type_var = tk.StringVar(value=contract[0] if contract else "CDD")
        salary_type_var = tk.StringVar(value=employee[19] if employee[19] else "في الساعة")
        entries = {}
        contract_entries = {}

        # Champs employé existants
        fields = [
            ("Matricule", "matricule", employee[0], True, ttk.Entry),
            ("Nom", "nom", employee[1], False, ttk.Entry),
            ("Prénom", "prenom", employee[2], False, ttk.Entry),
            ("Date Naissance (JJ/MM/AAAA)", "date_naissance", employee[4], False, DateEntry),
            ("Lieu Naissance", "lieu_naissance", employee[5], False, ttk.Entry),
            ("Adresse", "adresse", employee[6], False, ttk.Entry),
            ("Ville", "ville", employee[7], False, ttk.Entry),
            ("CIN", "cin", employee[8], False, ttk.Entry),
            ("Date CIN (JJ/MM/AAAA)", "date_cin", employee[9], False, DateEntry),
            ("Lieu CIN", "lieu_cin", employee[10], False, ttk.Entry),
            ("Poste", "poste", employee[11], False, ttk.Entry),
            ("Email", "email", employee[12], False, ttk.Entry),
            ("Téléphone", "telephone", employee[13], False, ttk.Entry),
        ]

        for i, (label, field, value, disabled, widget_type) in enumerate(fields):
            ttk.Label(employee_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5,
                                                                              sticky=tk.E)
            if widget_type == DateEntry:
                entry = widget_type(employee_frame, bootstyle="primary", dateformat="%d/%m/%Y")
                if value:
                    entry.entry.delete(0, tk.END)
                    entry.entry.insert(0, value)
            else:
                entry = widget_type(employee_frame, bootstyle="primary")
                if value:
                    entry.insert(0, value)
            if disabled:
                entry.config(state='disabled')
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=tk.EW)
            entries[field] = entry

        # Genre
        ttk.Label(employee_frame, text="Genre*", font=('Segoe UI', 10)).grid(row=len(fields), column=0, sticky=tk.E,
                                                                             padx=5, pady=5)
        genre_frame = ttk.Frame(employee_frame)
        genre_frame.grid(row=len(fields), column=1, sticky=tk.W)
        ttk.Radiobutton(genre_frame, text="السيدة", variable=genre_var, value="السيدة",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(genre_frame, text="الانسة", variable=genre_var, value="الانسة",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(genre_frame, text="السيد", variable=genre_var, value="السيد",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)

        # Champs contrat existants
        contract_fields = [
            ("Type de Contrat", "type_contrat", contract[0] if contract else "CDD", False, None),
            ("Date Début (JJ/MM/AAAA)*", "date_debut", employee[15], False, DateEntry),
            ("Date Fin (JJ/MM/AAAA)", "date_fin", employee[16], contract_type_var.get() == "CDI", DateEntry),
            ("Salaire Base*", "salaire", str(employee[17]) if employee[17] is not None else "", False, ttk.Entry),
            ("Prime*", "prime", str(employee[18]) if employee[18] is not None else "", False, ttk.Entry)
        ]

        for i, (label, field, value, disabled, widget_type) in enumerate(contract_fields):
            ttk.Label(contract_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5,
                                                                              sticky=tk.E)
            if field == "type_contrat":
                frame = ttk.Frame(contract_frame)
                frame.grid(row=i, column=1, sticky=tk.W)
                ttk.Radiobutton(frame, text="CDD", variable=contract_type_var, value="CDD",
                                bootstyle="primary-toolbutton",
                                command=lambda: contract_entries['date_fin'].entry.config(state=tk.NORMAL)).pack(
                    side=tk.LEFT, padx=5)
                ttk.Radiobutton(frame, text="CDI", variable=contract_type_var, value="CDI",
                                bootstyle="primary-toolbutton",
                                command=lambda: contract_entries['date_fin'].entry.config(state=tk.DISABLED)).pack(
                    side=tk.LEFT, padx=5)
            else:
                if widget_type == DateEntry:
                    entry = widget_type(contract_frame, bootstyle="primary", dateformat="%d/%m/%Y")
                    if value:
                        entry.entry.delete(0, tk.END)
                        entry.entry.insert(0, value)
                    if disabled:
                        entry.entry.config(state='disabled')
                else:
                    entry = widget_type(contract_frame, bootstyle="primary")
                    if value:
                        entry.insert(0, value)
                entry.grid(row=i, column=1, padx=5, pady=5, sticky=tk.EW)
                contract_entries[field] = entry

        # Type de salaire
        ttk.Label(contract_frame, text="Type de Salaire*", font=('Segoe UI', 10)).grid(row=len(contract_fields),
                                                                                       column=0, padx=5, pady=5,
                                                                                       sticky=tk.E)
        salary_type_frame = ttk.Frame(contract_frame)
        salary_type_frame.grid(row=len(contract_fields), column=1, sticky=tk.W)
        ttk.Radiobutton(salary_type_frame, text="الساعة", variable=salary_type_var, value="في الساعة",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)
        ttk.Radiobutton(salary_type_frame, text="الشهر", variable=salary_type_var, value="في الشهر",
                        bootstyle="primary-toolbutton").pack(side=tk.LEFT, padx=5)

        # NOUVEAUX CHAMPS: Dcon, duree, atelier, nbre_eche, Fperiode
        additional_fields = [
            ("Dcon", "dcon", employee[20], False, ttk.Entry),
            ("Durée", "duree", employee[21], False, ttk.Entry),
            ("Atelier", "atelier", employee[22], False, ttk.Entry),
            ("Nb Échéances", "nbre_eche", employee[23], False, ttk.Entry),
            ("Fperiode", "fperiode", employee[24], False, DateEntry)
        ]

        start_row = len(contract_fields) + 2
        for i, (label, field, value, disabled, widget_type) in enumerate(additional_fields):
            ttk.Label(contract_frame, text=label, font=('Segoe UI', 10)).grid(row=start_row + i, column=0, padx=5,
                                                                              pady=5, sticky=tk.E)

            if widget_type == DateEntry:
                entry = widget_type(contract_frame, bootstyle="primary", dateformat="%d/%m/%Y")
                if value:
                    entry.entry.delete(0, tk.END)
                    entry.entry.insert(0, value)
            else:
                entry = widget_type(contract_frame, bootstyle="primary")
                if value:
                    entry.insert(0, value)

            entry.grid(row=start_row + i, column=1, padx=5, pady=5, sticky=tk.EW)
            contract_entries[field] = entry

        # Boutons
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(fill=tk.X, pady=10)
        ttk.Button(button_frame, text="Enregistrer",
                   command=lambda: self.save_employee_and_contract_changes(matricule, entries, contract_entries,
                                                                           genre_var, contract_type_var,
                                                                           salary_type_var, top),
                   bootstyle=SUCCESS).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Annuler", command=top.destroy, bootstyle=WARNING).pack(side=tk.LEFT, padx=5)

        employee_frame.columnconfigure(1, weight=1)
        contract_frame.columnconfigure(1, weight=1)

        def on_mousewheel(event):
            try:
                delta = 0
                if event.num == 4:
                    delta = -1
                elif event.num == 5:
                    delta = 1
                elif event.delta:
                    delta = -1 * (event.delta // 120)
                if delta:
                    canvas.yview_scroll(int(delta), "units")
            except tk.TclError:
                pass

        top.bind("<MouseWheel>", on_mousewheel)
        top.bind("<Button-4>", on_mousewheel)
        top.bind("<Button-5>", on_mousewheel)

        def on_destroy():
            top.unbind("<MouseWheel>")
            top.unbind("<Button-4>")
            top.unbind("<Button-5>")
            top.destroy()

        top.protocol("WM_DELETE_WINDOW", on_destroy)

    # 4. Modifier save_employee_and_contract_changes
    def save_employee_and_contract_changes(self, matricule, entries, contract_entries, genre_var, contract_type_var,
                                           salary_type_var, top):
        """Save changes to employee and contract details, updating only modified fields."""
        try:
            with self.conn.cursor() as cursor:
                cursor.execute('''
                    SELECT matricule, nom, prenom, genre, date_naissance, lieu_naissance,
                           adresse, ville, cin, date_cin, lieu_cin, poste, email, telephone,
                           type_contrat, date_debut, date_fin, salaire_base, prime, salary_type,
                           dcon, duree, atelier, nbre_eche, fperiode
                    FROM employees
                    WHERE matricule = %s
                ''', (matricule,))
                original_employee = cursor.fetchone()

            if not original_employee:
                Messagebox.show_error("Employé non trouvé", "Erreur", parent=top)
                top.destroy()
                return

            # Récupérer les valeurs des nouveaux champs
            dcon = self.get_widget_value(contract_entries.get("dcon", ""))
            duree = self.get_widget_value(contract_entries.get("duree", ""))
            atelier = self.get_widget_value(contract_entries.get("atelier", ""))
            nbre_eche = self.get_widget_value(contract_entries.get("nbre_eche", ""))
            fperiode = self.get_widget_value(contract_entries.get("fperiode", ""))

            # Convertir nbre_eche en entier si possible
            try:
                nbre_eche = int(nbre_eche) if nbre_eche else None
            except ValueError:
                nbre_eche = None

            employee_data = {
                "matricule": matricule,
                "nom": self.get_widget_value(entries["nom"]) or original_employee[1],
                "prenom": self.get_widget_value(entries["prenom"]) or original_employee[2],
                "genre": genre_var.get() or original_employee[3],
                "date_naissance": self.get_widget_value(entries["date_naissance"]) or original_employee[4],
                "lieu_naissance": self.get_widget_value(entries["lieu_naissance"]) or original_employee[5],
                "adresse": self.get_widget_value(entries["adresse"]) or original_employee[6],
                "ville": self.get_widget_value(entries["ville"]) or original_employee[7],
                "cin": self.get_widget_value(entries["cin"]) or original_employee[8],
                "date_cin": self.get_widget_value(entries["date_cin"]) or original_employee[9],
                "lieu_cin": self.get_widget_value(entries["lieu_cin"]) or original_employee[10],
                "poste": self.get_widget_value(entries["poste"]) or original_employee[11],
                "email": self.get_widget_value(entries["email"]) or original_employee[12],
                "telephone": self.get_widget_value(entries["telephone"]) or original_employee[13],
                "type_contrat": contract_type_var.get() or original_employee[14],
                "date_debut": self.get_widget_value(contract_entries["date_debut"]) or original_employee[15],
                "date_fin": self.get_widget_value(
                    contract_entries["date_fin"]) if contract_type_var.get() == "CDD" else None,
                "salaire_base": float(self.get_widget_value(contract_entries["salaire"])) if self.get_widget_value(
                    contract_entries["salaire"]) else original_employee[17],
                "prime": float(self.get_widget_value(contract_entries["prime"])) if self.get_widget_value(
                    contract_entries["prime"]) else original_employee[18],
                "salary_type": salary_type_var.get() or original_employee[19],
                "dcon": dcon or original_employee[20],
                "duree": duree or original_employee[21],
                "atelier": atelier or original_employee[22],
                "nbre_eche": nbre_eche if nbre_eche is not None else original_employee[23],
                "fperiode": fperiode or original_employee[24]
            }

            # Générer le texte du contrat
            contrat_text = self.generate_contract_from_data(employee_data)

            # Mettre à jour la base de données
            with self.conn.cursor() as cursor:
                cursor.execute('''
                    UPDATE employees 
                    SET nom=%s, prenom=%s, genre=%s, date_naissance=%s, lieu_naissance=%s,
                        adresse=%s, ville=%s, cin=%s, date_cin=%s, lieu_cin=%s, poste=%s,
                        email=%s, telephone=%s, type_contrat=%s, date_debut=%s, date_fin=%s,
                        salaire_base=%s, prime=%s, salary_type=%s, dcon=%s, duree=%s,
                        atelier=%s, nbre_eche=%s, fperiode=%s
                    WHERE matricule=%s
                ''', (
                    employee_data['nom'],
                    employee_data['prenom'],
                    employee_data['genre'],
                    employee_data['date_naissance'],
                    employee_data['lieu_naissance'],
                    employee_data['adresse'],
                    employee_data['ville'],
                    employee_data['cin'],
                    employee_data['date_cin'],
                    employee_data['lieu_cin'],
                    employee_data['poste'],
                    employee_data['email'],
                    employee_data['telephone'],
                    employee_data['type_contrat'],
                    employee_data['date_debut'],
                    employee_data['date_fin'],
                    employee_data['salaire_base'],
                    employee_data['prime'],
                    employee_data['salary_type'],
                    employee_data['dcon'],
                    employee_data['duree'],
                    employee_data['atelier'],
                    employee_data['nbre_eche'],
                    employee_data['fperiode'],
                    matricule
                ))

                cursor.execute('''
                    INSERT INTO contrats (matricule, type_contrat, date_creation, texte_contrat)
                    VALUES (%s, %s, NOW(), %s)
                    ON DUPLICATE KEY UPDATE 
                        type_contrat=VALUES(type_contrat), 
                        date_creation=NOW(), 
                        texte_contrat=VALUES(texte_contrat)
                ''', (matricule, employee_data["type_contrat"], contrat_text))

                self.conn.commit()

            self.load_data()
            self.status_var.set(f"Employé {matricule} mis à jour avec succès")
            Messagebox.show_info("Employé mis à jour avec succès", "Succès", parent=top)
            top.destroy()

        except Exception as e:
            self.conn.rollback()
            Messagebox.show_error(f"Erreur lors de la sauvegarde: {str(e)}", "Erreur", parent=top)

    # Dans la méthode load_alert_table()
    def load_alert_table(self):
        """Load the alerts table with expiring contracts including CDI."""
        try:
            self.alert_table.delete_rows()
            today = datetime.datetime.now().date()

            with self.conn.cursor() as cursor:
                cursor.execute('''
                    SELECT e.matricule,
                           e.nom,
                           e.prenom,
                           COALESCE(e.date_fin, e.fperiode) as end_date,
                           e.atelier,
                           e.nbre_eche,
                           e.type_contrat,
                           e.fperiode
                    FROM employees e
                    WHERE (e.type_contrat = 'CDD' OR e.type_contrat = 'CDI')
                      AND COALESCE(e.date_fin, e.fperiode) IS NOT NULL
                      AND COALESCE(e.date_fin, e.fperiode) != ''
                      AND STR_TO_DATE(COALESCE(e.date_fin, e.fperiode), '%%Y-%%m-%%d') >= %s
                      AND DATEDIFF(STR_TO_DATE(COALESCE(e.date_fin, e.fperiode), '%%Y-%%m-%%d'), %s) BETWEEN 0 AND 30
                    ORDER BY STR_TO_DATE(COALESCE(e.date_fin, e.fperiode), '%%Y-%%m-%%d')
                ''', (today, today))

                rows = cursor.fetchall()

                for i, row in enumerate(rows):
                    matricule, nom, prenom, end_date, atelier, nbre_eche, type_contrat, fperiode = row

                    try:
                        end_date_obj = datetime.datetime.strptime(str(end_date), "%Y-%m-%d").date()
                        jours_restants = (end_date_obj - today).days
                        end_date_str = end_date_obj.strftime('%d/%m/%Y')
                    except (ValueError, TypeError):
                        jours_restants = "N/A"
                        end_date_str = str(end_date)

                    fperiode_str = fperiode if fperiode else "N/A"

                    self.alert_table.insert_row(
                        values=[
                            matricule,
                            nom or "N/A",
                            prenom or "N/A",
                            end_date_str,
                            str(jours_restants),
                            atelier or "N/A",
                            str(nbre_eche) if nbre_eche is not None else "0",
                            type_contrat or "N/A",
                            fperiode_str
                        ]
                    )

            self.alert_table.load_table_data()
            self.status_var.set(f"{len(rows)} contrats expirant bientôt")

        except Exception as e:
            logging.error(f"Erreur lors du chargement des alertes: {str(e)}")
            Messagebox.show_error(f"Erreur lors du chargement des alertes: {str(e)}", "Erreur")
            self.status_var.set("Erreur de chargement des alertes")

    # Dans la méthode background_alert_service()
    from plyer import notification

    from plyer import notification
    import datetime
    import logging

    from plyer import notification
    import datetime
    import logging

    def background_alert_service(self):
        """Service en arrière-plan pour vérifier les contrats expirants (CDD et CDI)"""
        try:
            today = datetime.datetime.now().date()

            with self.conn.cursor() as cursor:
                cursor.execute('''
                    SELECT e.matricule,
                           e.nom,
                           e.prenom,
                           COALESCE(e.date_fin, e.fperiode) as end_date,
                           e.type_contrat,
                           e.fperiode
                    FROM employees e
                    WHERE (e.type_contrat = 'CDD' OR e.type_contrat = 'CDI')
                      AND COALESCE(e.date_fin, e.fperiode) IS NOT NULL
                      AND COALESCE(e.date_fin, e.fperiode) != ''
                      AND STR_TO_DATE(COALESCE(e.date_fin, e.fperiode), '%%Y-%%m-%%d') >= %s
                      AND DATEDIFF(STR_TO_DATE(COALESCE(e.date_fin, e.fperiode), '%%Y-%%m-%%d'), %s) BETWEEN 0 AND 30
                ''', (today, today))

                expiring_contracts = cursor.fetchall()

            # Comptage réel pour la notification
            total_expiring_contracts = len(expiring_contracts)

            new_contracts = []
            for contract in expiring_contracts:
                matricule, nom, prenom, end_date, contract_type, fperiode = contract

                try:
                    end_date_obj = datetime.datetime.strptime(str(end_date), "%Y-%m-%d").date()
                    end_date_str = end_date_obj.strftime('%Y-%m-%d')
                except (ValueError, TypeError):
                    end_date_str = str(end_date)

                contract_key = f"{matricule}_{end_date_str}"

                # On considère seulement les contrats jamais alertés dans les dernières 24h
                if (contract_key not in self.alerted_contracts or
                        (datetime.datetime.now() - self.alerted_contracts[contract_key][
                            'last_alerted']).total_seconds() >= 24 * 3600):

                    try:
                        end_date_obj = datetime.datetime.strptime(str(end_date), "%Y-%m-%d").date()
                        jours_restants = (end_date_obj - today).days
                    except (ValueError, TypeError):
                        jours_restants = "N/A"

                    new_contracts.append((matricule, nom, prenom, end_date, jours_restants, contract_type, fperiode))
                    self.alerted_contracts[contract_key] = {
                        'end_date': end_date,
                        'last_alerted': datetime.datetime.now()
                    }

            # Mise à jour UI pour les nouveaux contrats seulement
            if new_contracts:
                self.show_contract_alerts(new_contracts)
                self.load_alert_table()

            # --- Notification système toutes les minutes ---
            try:
                notification.notify(
                    title="Alertes Contrats",
                    message=f"{total_expiring_contracts} contrat(s) expirant dans les 30 prochains jours !",
                    timeout=10
                )
            except Exception as e:
                logging.error(f"Erreur notification: {str(e)}")
            # -----------------------------------------------

        except Exception as e:
            logging.error(f"Erreur dans background_alert_service: {str(e)}")
        finally:
            # Relance la fonction après l'intervalle défini (en ms)
            self.alert_timer = self.root.after(self.check_interval, self.background_alert_service)

    # Dans la méthode show_contract_alerts()
    def show_contract_alerts(self, contracts):
        """Affiche les alertes de contrat"""
        message = "⚠️ ALERTE : Contrats expirant bientôt ⚠️\n\n"
        for contract in contracts:
            matricule, nom, prenom, end_date, jours_restants, contract_type, fperiode = contract
            message += f"• {nom} {prenom} (Matricule: {matricule}, Type: {contract_type}) - "
            message += f"Expire le {end_date} (dans {jours_restants} jours)"
            if fperiode:
                message += f", Fperiode: {fperiode}"
            message += "\n"

        self.play_alert_sound()
        Messagebox.show_warning(message, "Alerte Contrats", parent=self.root)
        self.stop_alert_sound()
        self.status_var.set(f"⚠ {len(contracts)} nouveaux contrats expirent bientôt")
#####################################################################################################


    def __del__(self):
        self.stop_alert_timer()
        self.stop_alert_sound()
        if hasattr(self, 'conn'):
            try:
                self.conn.close()
            except:
                pass

    def get_widget_value(self, widget):
        """Obtient la valeur de n'importe quel widget"""
        if isinstance(widget, (ttk.Entry, ttk.Combobox)):
            return widget.get().strip()
        elif isinstance(widget, DateEntry):
            return widget.entry.get().strip()
        elif hasattr(widget, 'get'):
            return widget.get().strip()
        return ""

    def create_contract_doc(self, matricule, texte=None):
        with self.conn.cursor() as cursor:
            cursor.execute('''
                           SELECT e.matricule,
                                  e.nom,
                                  e.prenom,
                                  e.genre,
                                  e.date_naissance,
                                  e.lieu_naissance,
                                  e.ville,
                                  e.cin,
                                  e.date_cin,
                                  e.lieu_cin,
                                  e.poste,
                                  e.type_contrat,
                                  e.date_debut,
                                  e.date_fin,
                                  e.salaire_base,
                                  e.prime,
                                  e.salary_type,
                                  e.adresse,
                                  e.date_embauche,
                                  e.code_postal,
                                  e.dcon,
                                  e.duree,
                                  e.atelier,
                                  e.nbre_eche,
                                  e.fperiode,
                                  e.degre_polyvalence,
                                  c.texte_contrat
                           FROM employees e
                                    LEFT JOIN contrats c ON e.matricule = c.matricule
                           WHERE e.matricule = %s
                           ORDER BY c.date_creation DESC LIMIT 1
                           ''', (matricule,))
            result = cursor.fetchone()

            if not result:
                print("Employé non trouvé", "Erreur")
                return None

            employee_data = {
                'matricule': result[0] or "غير محدد",
                'nom': result[1] or "غير محدد",
                'prenom': result[2] or "غير محدد",
                'genre': result[3] or "غير محدد",
                'date_naissance': result[4] or None,
                'lieu_naissance': result[5] or "غير محدد",
                'ville': result[6] or "غير محدد",
                'cin': result[7] or "غير محدد",
                'date_cin': result[8] or None,
                'lieu_cin': result[9] or "غير محدد",
                'poste': result[10] or "غير محدد",
                'type_contrat': result[11] or "غير محدد",
                'date_debut': result[12] or None,
                'date_fin': result[13] or None,
                'salaire_base': float(result[14] or 0),
                'prime': float(result[15] or 0),
                'salary_type': "الساعة" if result[16] == "في الساعة" else "الشهر",
                'adresse': result[17] or "غير محدد",
                'date_embauche': result[18] or None,
                'code_postal': result[19] or "غير محدد",
                'dcon': result[20] or "غير محدد",
                'duree': result[21] or "غير محدد",
                'atelier': result[22] or "غير محدد",
                'nbre_eche': result[23] or 0,
                'fperiode': result[24] or "غير محدد",
                'degre_polyvalence': result[25] or "غير محدد"
            }
            texte_contrat = result[26] if result[26] else texte

        if not texte_contrat:
            texte_contrat = self.generate_contract_from_data(employee_data)

        doc = Document()
        section = doc.sections[0]
        section.left_margin = section.right_margin = Inches(8 / 25.4)
        section.top_margin = section.bottom_margin = Inches(5 / 25.4)
        section.is_right_to_left = True

        # ======================================================================
        # CORPS DU DOCUMENT (sans en-tête)
        # ======================================================================
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # Matricule de l'employé
        matricule_para = doc.add_paragraph()
        matricule_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        matricule_run = matricule_para.add_run(f"Matricule: {employee_data['matricule']}")
        matricule_run.bold = True
        matricule_run.font.name = "Arial"
        matricule_run.font.size = Pt(9)

        # Espace avant le texte du contrat
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # Ajout du texte du contrat
        for paragraph in texte_contrat.strip().split('\n'):
            if paragraph.strip():
                body_para = doc.add_paragraph()
                body_para.paragraph_format.space_after = Pt(4)
                body_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                body_run = body_para.add_run(paragraph.strip())
                body_run.font.name = "Arial"
                body_run.font.size = Pt(9)
                body_run.font.rtl = True

        return doc

    def select_contract_template(self, genre, type_contrat):
        # Normalize inputs
        genre = (genre or "").strip()
        type_contrat = (type_contrat or "").strip().upper()

        # Log for debugging
        print(f"Selecting template with genre: '{genre}', type_contrat: '{type_contrat}'")

        if type_contrat == "CDD":
            if genre in ["السيدة","الانسة","الأنسة"]:
                return self.CDD_FEMININ
            elif genre == "السيد":
                return self.CDD_MASCULIN
        elif type_contrat == "CDI":
            if genre in ["السيدة","الانسة","الأنسة"]:
                return self.CDI_FEMININ
            elif genre == "السيد":
                return self.CDI_MASCULIN

        # Log warning for invalid inputs
        print(f"Warning: No template found for genre='{genre}', type_contrat='{type_contrat}'")
        Messagebox.show_warning(
            f"Données invalides: genre='{genre}', type_contrat='{type_contrat}'. Veuillez vérifier les données de l'employé.",
            "Avertissement"

        )
        return None

    import re
    import datetime
    from dateutil.relativedelta import relativedelta

    def generate_contract_from_data(self, employee_data):
        def safe_date_format(date_str, default="غير محدد"):
            if not date_str:
                print(f"Date is None or empty: {date_str}")
                return default
            try:
                if isinstance(date_str, str):
                    if re.match(r'^\d{4}-\d{2}-\d{2}$', date_str):
                        date_obj = datetime.datetime.strptime(date_str, "%Y-%m-%d")
                        return date_obj.strftime("%d/%m/%Y")
                    elif re.match(r'^\d{2}/\d{2}/\d{4}$', date_str):
                        return date_str
                    for fmt in ["%Y-%m-%d", "%d/%m/%Y"]:
                        try:
                            date_obj = datetime.datetime.strptime(date_str, fmt)
                            return date_obj.strftime("%d/%m/%Y")
                        except ValueError:
                            continue
                    print(f"Failed to parse date: {date_str}")
                    return default
                elif isinstance(date_str, datetime.date):
                    return date_str.strftime("%d/%m/%Y")
                print(f"Unexpected date type: {type(date_str)}")
                return default
            except Exception as e:
                print(f"Error parsing date {date_str}: {str(e)}")
                return default

        # Préparer les données de l'employé
        data = {
            'Titre': employee_data.get('genre', "غير محدد"),
            'Prénom': employee_data.get('prenom', "غير محدد"),
            'Nom': employee_data.get('nom', "غير محدد"),
            'DNAIS': safe_date_format(employee_data.get('date_naissance')),
            'LNAIS': employee_data.get('lieu_naissance', 'غير محدد'),
            'Ville': employee_data.get('adresse', 'غير محدد'),
            'NCIN': employee_data.get('cin', 'غير محدد'),
            'DCIN': safe_date_format(employee_data.get('date_cin')),
            'LCIN': employee_data.get('lieu_cin', 'غير محدد'),
            'Poste': employee_data.get('poste', 'غير محدد'),
            'DPERIODE': safe_date_format(employee_data.get('date_debut')),
            'FPERIODE': safe_date_format(employee_data.get('date_fin')),
            'SBASE': str(employee_data.get('salaire_base', 0)),
            'PRIME': str(employee_data.get('prime', 0)),
            'MPAIE': "الساعة" if employee_data.get('salary_type') == "في الساعة" else "الشهر",
            'DATE_CONTRAT': datetime.datetime.now().strftime('%d/%m/%Y'),
            'DUREE': employee_data.get('duree', 'غير محددة')
        }

        # Calculer la durée exacte pour les CDD
        if (employee_data.get('type_contrat') == "CDD" and
                data.get('DPERIODE') != "غير محدد" and
                data.get('FPERIODE') != "غير محدد"):
            try:
                debut = datetime.datetime.strptime(data['DPERIODE'], "%d/%m/%Y")
                fin = datetime.datetime.strptime(data['FPERIODE'], "%d/%m/%Y")
                rd = relativedelta(fin, debut)
                months_total = rd.years * 12 + rd.months
                days = rd.days
                data['DUREE'] = f"{months_total} شهرا" + (f" و {days} يوما" if days else "")
            except (ValueError, TypeError) as e:
                print(f"Error calculating DUREE: {str(e)}")
                data['DUREE'] = "غير محددة"

        # Sélectionner le bon template
        template = self.select_contract_template(
            employee_data.get('genre'),
            employee_data.get('type_contrat')
        )

        if not template:
            raise ValueError("No contract template found for the given genre and contract type")

        # Remplacer les placeholders
        contrat_text = template
        for placeholder, value in data.items():
            contrat_text = contrat_text.replace(f'{{{{{placeholder}}}}}', str(value))

        return contrat_text

    def save_and_generate(self):
        try:
            # Validation des champs obligatoires
            required_fields = {
                'matricule': self.get_widget_value(self.entries['matricule']),
                'nom': self.get_widget_value(self.entries['nom']),
                'prenom': self.get_widget_value(self.entries['prenom']),
                'date_debut': self.get_widget_value(self.contract_entries['date_debut']),
                'salaire': self.get_widget_value(self.contract_entries['salaire']),
                'prime': self.get_widget_value(self.contract_entries['prime'])
            }

            for field, value in required_fields.items():
                if not value:
                    Messagebox.show_warning(f"Le champ {field} est obligatoire", "Erreur")
                    return

            matricule = required_fields['matricule']

            # Vérifier si le matricule existe déjà
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT 1 FROM employees WHERE matricule = %s", (matricule,))
                if cursor.fetchone():
                    Messagebox.show_warning(f"Le matricule {matricule} existe déjà", "Erreur")
                    return

            # Convertir les dates au format MySQL (YYYY-MM-DD)
            def convert_to_mysql_date(date_str):
                try:
                    if not date_str:
                        return None
                    # Convertir de DD/MM/YYYY vers YYYY-MM-DD
                    date_obj = datetime.datetime.strptime(date_str, "%d/%m/%Y")
                    return date_obj.strftime("%Y-%m-%d")
                except ValueError:
                    return None

            # Préparer les données pour l'insertion
            employee_data = {
                'matricule': matricule,
                'nom': required_fields['nom'],
                'prenom': required_fields['prenom'],
                'genre': self.variables["genre"].get(),
                'date_naissance': convert_to_mysql_date(self.get_widget_value(self.entries['date_naissance'])),
                'lieu_naissance': self.get_widget_value(self.entries['lieu_naissance']) or None,
                'adresse': self.get_widget_value(self.entries['adresse']) or None,
                'ville': self.get_widget_value(self.entries['ville']) or "المحرس",
                'cin': self.get_widget_value(self.entries['cin']) or None,
                'date_cin': convert_to_mysql_date(self.get_widget_value(self.entries['date_cin'])),
                'lieu_cin': self.get_widget_value(self.entries['lieu_cin']) or "تونس",
                'poste': self.get_widget_value(self.entries['poste']) or None,
                'email': self.get_widget_value(self.entries['email']) or None,
                'telephone': self.get_widget_value(self.entries['telephone']) or None,
                'type_contrat': self.variables["contract_type"].get(),
                'date_debut': convert_to_mysql_date(required_fields['date_debut']),
                'date_fin': convert_to_mysql_date(self.get_widget_value(self.contract_entries['date_fin'])) if
                self.variables["contract_type"].get() == "CDD" else None,
                'salaire_base': float(required_fields['salaire']),
                'prime': float(required_fields['prime']),
                'salary_type': self.variables["salary_type"].get(),
                'duree': self.get_widget_value(self.entries['duree']) or None,
                'atelier': self.get_widget_value(self.entries['atelier']) or None,
                'nbre_eche': self.get_widget_value(self.entries['nbre_eche']) or None,
                'fperiode': convert_to_mysql_date(self.get_widget_value(self.entries['fperiode'])) or None
            }

            # Générer le texte du contrat
            contrat_text = self.generate_contract_from_data(employee_data)

            # Insertion dans la base de données
            with self.conn.cursor() as cursor:
                # Insertion dans la table employees
                cursor.execute('''
                    INSERT INTO employees (
                        matricule, nom, prenom, genre, date_naissance, lieu_naissance,
                        adresse, ville, cin, date_cin, lieu_cin, poste, email, telephone,
                        type_contrat, date_debut, date_fin, salaire_base, prime, salary_type,
                        duree, atelier, nbre_eche, fperiode
                    ) VALUES (
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                    )
                ''', (
                    employee_data['matricule'],
                    employee_data['nom'],
                    employee_data['prenom'],
                    employee_data['genre'],
                    employee_data['date_naissance'],
                    employee_data['lieu_naissance'],
                    employee_data['adresse'],
                    employee_data['ville'],
                    employee_data['cin'],
                    employee_data['date_cin'],
                    employee_data['lieu_cin'],
                    employee_data['poste'],
                    employee_data['email'],
                    employee_data['telephone'],
                    employee_data['type_contrat'],
                    employee_data['date_debut'],
                    employee_data['date_fin'],
                    employee_data['salaire_base'],
                    employee_data['prime'],
                    employee_data['salary_type'],
                    employee_data['duree'],
                    employee_data['atelier'],
                    employee_data['nbre_eche'],
                    employee_data['fperiode']
                ))

                # Insertion dans la table contrats
                cursor.execute('''
                    INSERT INTO contrats (matricule, type_contrat, date_creation, texte_contrat)
                    VALUES (%s, %s, NOW(), %s)
                ''', (
                    employee_data['matricule'],
                    employee_data['type_contrat'],
                    contrat_text
                ))

                self.conn.commit()

            # Mettre à jour l'interface
            self.current_employee = employee_data
            self.contract_text.config(state=tk.NORMAL)
            self.contract_text.delete(1.0, tk.END)
            self.contract_text.insert(tk.END, contrat_text, 'rtl')
            self.contract_text.config(state=tk.DISABLED)
            self.notebook.select(2)

            # Actualiser les données
            self.load_data()
            self.load_matricules()
            self.clear_form()

            self.status_var.set(f"Contrat {employee_data['type_contrat']} enregistré (N°{matricule})")
            Messagebox.show_info(f"Contrat {matricule} enregistré avec succès", "Succès")

        except pymysql.MySQLError as e:
            self.conn.rollback()
            error_msg = f"Erreur MySQL: {str(e)}"
            Messagebox.showerror("Erreur base de données", error_msg)

        except ValueError as e:
            error_msg = f"Erreur de valeur: {str(e)}"
            Messagebox.showerror("Erreur de donnée", error_msg)

        except Exception as e:
            error_msg = f"Erreur inattendue: {str(e)}"
            Messagebox.showerror("Erreur", error_msg)




    def setup_system_tray_icon(self):
        """Configurer l'icône de la barre système pour un accès rapide"""
        try:
            import pystray
            from PIL import Image
            import threading

            # Créer une image pour l'icône
            image = Image.new('RGB', (64, 64), color='white')

            # Créer le menu
            menu = pystray.Menu(
                pystray.MenuItem("Ouvrir", self.show_application),
                pystray.MenuItem("Vérifier maintenant", self.check_contracts_now),
                pystray.MenuItem("Quitter", self.quit_application)
            )

            # Créer l'icône
            self.tray_icon = pystray.Icon("contrat_app", image, "Gestion des Contrats", menu)

            # Démarrer l'icône dans un thread séparé
            self.tray_thread = threading.Thread(target=self.tray_icon.run, daemon=True)
            self.tray_thread.start()

        except ImportError:
            print("Les bibliothèques pystray et PIL sont nécessaires pour la barre système")

    def show_application(self, icon, item):
        """Afficher l'application"""
        self.root.deiconify()
        self.root.state('zoomed')

    def check_contracts_now(self, icon, item):
        """Vérifier les contrats immédiatement"""
        self.check_expiring_contracts(force_notification=True)

    def quit_application(self, icon, item):
        """Quitter l'application"""
        self.root.quit()

    def on_close(self):
        """Gérer la fermeture de l'application"""
        if hasattr(self, 'tray_icon'):
            self.tray_icon.visible = False
        self.root.destroy()






    def load_employee_table(self):
        try:
            with self.conn.cursor() as cursor:
                cursor.execute('''
                               SELECT matricule,
                                      nom,
                                      prenom,
                                      genre,
                                      date_naissance,
                                      lieu_naissance,
                                      adresse,
                                      ville,
                                      cin,
                                      date_cin,
                                      lieu_cin,
                                      poste,
                                      email,
                                      telephone,
                                      type_contrat,
                                      date_debut,
                                      date_fin,
                                      salaire_base,
                                      prime,
                                      salary_type,
                                      atelier,
                                      nbre_eche
                               FROM employees
                               ''')
                rows = cursor.fetchall()

            # Fonction pour formater les dates
            def format_date_for_display(date_str):
                try:
                    if not date_str:
                        return 'N/A'
                    # Si la date est déjà au bon format, la retourner telle quelle
                    if isinstance(date_str, str) and '/' in date_str:
                        return date_str
                    # Sinon, convertir du format YYYY-MM-DD vers DD/MM/YYYY
                    date_obj = datetime.datetime.strptime(str(date_str), "%Y-%m-%d")
                    return date_obj.strftime("%d/%m/%Y")
                except (ValueError, TypeError):
                    return 'N/A'

            # Préparer les données avec les dates formatées
            formatted_rows = []
            for row in rows:
                formatted_row = list(row)
                # Formater les dates (index 4, 9, 15, 16)
                formatted_row[4] = format_date_for_display(row[4])  # date_naissance
                formatted_row[9] = format_date_for_display(row[9])  # date_cin
                formatted_row[15] = format_date_for_display(row[15])  # date_debut
                formatted_row[16] = format_date_for_display(row[16])  # date_fin
                formatted_rows.append(formatted_row)

            self.update_table_data(formatted_rows)

        except Exception as e:
            print(f"Erreur lors du chargement des employés: {str(e)}", "Erreur")
            self.status_var.set("Erreur de chargement")

    def generate_and_show_contract(self):
        try:
            if not self.current_employee:
                Messagebox.show_warning("Aucun employé sélectionné", "Attention")
                return

            # Vérifier si un contrat existe déjà pour cet employé
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT COUNT(*) FROM contrats WHERE matricule = %s",
                               (self.current_employee['matricule'],))
                if cursor.fetchone()[0] > 0:
                    # Si contrat existe déjà, demander confirmation pour regénérer
                    if not Messagebox.yesno("Un contrat existe déjà. Voulez-vous le regénérer ?",
                                            "Confirmation"):
                        return

            # Générer le nouveau contrat
            contrat_text = self.generate_contract_from_data(self.current_employee)

            with self.conn.cursor() as cursor:
                # Supprimer l'ancien contrat s'il existe
                cursor.execute("DELETE FROM contrats WHERE matricule = %s",
                               (self.current_employee['matricule'],))

                # Insérer le nouveau contrat
                cursor.execute("""
                               INSERT INTO contrats
                                   (matricule, date_creation, texte_contrat, type_contrat)
                               VALUES (%s, NOW(), %s, %s)
                               """, (
                                   self.current_employee['matricule'],
                                   contrat_text,
                                   self.current_employee['type_contrat']
                               ))
                self.conn.commit()

            # Afficher le contrat
            self.contract_text.config(state=tk.NORMAL)
            self.contract_text.delete(1.0, tk.END)
            self.contract_text.insert(tk.END, contrat_text, 'rtl')
            self.contract_text.config(state=tk.DISABLED)
            self.notebook.select(2)
            self.status_var.set("Contrat généré avec succès")

        except pymysql.Error as e:
            self.conn.rollback()
            print(f"Erreur MySQL: {str(e)}", "Erreur")
        except Exception as e:
            print(f"Erreur inattendue: {str(e)}", "Erreur")

    def update_table_data(self, rows):
        """Met à jour le tableau des employés avec les nouvelles données"""
        self.employee_table.delete_rows()
        today = datetime.datetime.now().date()

        for row in rows:
            matricule = row[0]

            # Vérifier si le contrat expire bientôt
            warning = ""
            date_fin = row[16]  # date_fin est à l'index 16
            if date_fin and isinstance(date_fin, str) and date_fin != "N/A":
                try:
                    # Essayer différents formats de date
                    try:
                        end_date = datetime.datetime.strptime(date_fin, "%d/%m/%Y").date()
                    except ValueError:
                        try:
                            end_date = datetime.datetime.strptime(date_fin, "%Y-%m-%d").date()
                        except ValueError:
                            end_date = None

                    if end_date:
                        days_left = (end_date - today).days
                        if 0 <= days_left <= 30:
                            warning = "⚠️ "
                except (ValueError, AttributeError):
                    pass

            # Préparation des données à afficher
            display_row = [
                row[0],  # matricule
                warning + (row[1] or ""),  # nom avec avertissement si besoin
                row[2] or "",  # prenom
                row[3] or "",  # genre
                row[4] or "N/A",  # date_naissance
                row[5] or "N/A",  # lieu_naissance
                row[6] or "N/A",  # adresse
                row[7] or "N/A",  # ville
                row[8] or "N/A",  # cin
                row[9] or "N/A",  # date_cin
                row[10] or "N/A",  # lieu_cin
                row[11] or "N/A",  # poste
                row[12] or "N/A",  # email
                row[13] or "N/A",  # telephone
                row[14] or "N/A",  # type_contrat
                row[15] or "N/A",  # date_debut
                row[16] or "N/A",  # date_fin
                f"{float(row[17]):.2f}" if row[17] is not None and row[17] != "N/A" else "N/A",  # salaire_base
                f"{float(row[18]):.2f}" if row[18] is not None and row[18] != "N/A" else "N/A",  # prime
                "Par heure" if row[19] == "hourly" else "Par mois" if row[19] == "monthly" else "N/A",  # salary_type
                row[20] or "N/A",  # atelier
                str(row[21]) if row[21] is not None else "0"  # nbre_eche
            ]

            # Ajout de la ligne dans le tableau
            self.employee_table.insert_row(values=display_row)

        # Mise à jour des statistiques
        self.update_summary()

    def get_table_row_safe(self, row_id):
        """Récupère une ligne du tableau de manière sécurisée"""
        try:
            # Méthode 1: Utiliser l'index numérique
            if isinstance(row_id, str) and row_id.startswith('I'):
                row_index = int(row_id[1:]) - 1
                if 0 <= row_index < len(self.employee_table.tablerows):
                    return self.employee_table.tablerows[row_index]

            # Méthode 2: Utiliser la méthode get_row si disponible
            if hasattr(self.employee_table, 'get_row'):
                return self.employee_table.get_row(row_id)

            return None
        except (ValueError, IndexError, AttributeError):
            return None

    def search_employee(self):
        matricule = self.variables["matricule"].get()
        if not matricule:
            Messagebox.show_warning("Veuillez saisir un matricule", "Attention")
            return

        try:
            with self.conn.cursor() as cursor:
                cursor.execute('''
                               SELECT matricule,
                                      nom,
                                      prenom,
                                      genre,
                                      date_naissance,
                                      lieu_naissance,
                                      adresse,
                                      ville,
                                      cin,
                                      date_cin,
                                      lieu_cin,
                                      poste,
                                      email,
                                      telephone,
                                      type_contrat,
                                      date_debut,
                                      date_fin,
                                      salaire_base,
                                      prime,
                                      salary_type
                               FROM employees
                               WHERE matricule = %s
                               ''', (matricule,))
                employee = cursor.fetchone()

            if employee:
                # Fonction pour formater les dates au format DD/MM/YYYY
                def format_date_for_display(date_str):
                    try:
                        if not date_str:
                            return 'N/A'
                        # Si la date est déjà au bon format, la retourner telle quelle
                        if isinstance(date_str, str) and '/' in date_str:
                            return date_str
                        # Sinon, convertir du format YYYY-MM-DD vers DD/MM/YYYY
                        date_obj = datetime.datetime.strptime(str(date_str), "%Y-%m-%d")
                        return date_obj.strftime("%d/%m/%Y")
                    except (ValueError, TypeError):
                        return 'N/A'

                # Conversion explicite en dictionnaire avec dates formatées
                employee_dict = {
                    'matricule': employee[0],
                    'nom': employee[1],
                    'prenom': employee[2],
                    'genre': employee[3],
                    'date_naissance': format_date_for_display(employee[4]),
                    'lieu_naissance': employee[5],
                    'adresse': employee[6],
                    'ville': employee[7],
                    'cin': employee[8],
                    'date_cin': format_date_for_display(employee[9]),
                    'lieu_cin': employee[10],
                    'poste': employee[11],
                    'email': employee[12],
                    'telephone': employee[13],
                    'type_contrat': employee[14],
                    'date_debut': format_date_for_display(employee[15]),
                    'date_fin': format_date_for_display(employee[16]),
                    'salaire_base': float(employee[17]) if employee[17] else 0.0,
                    'prime': float(employee[18]) if employee[18] else 0.0,
                    'salary_type': employee[19]
                }

                self.current_employee = employee_dict
                # Afficher les informations
                info_text = f"""Matricule: {employee[0]}
    Nom: {employee[1]} {employee[2]}
    Genre: {employee[3]}
    Date Naissance: {format_date_for_display(employee[4])} à {employee[5] or 'N/A'}
    Adresse: {employee[6] or 'N/A'}, {employee[7] or 'N/A'}
    CIN: {employee[8] or 'N/A'} (délivré le {format_date_for_display(employee[9])} à {employee[10] or 'N/A'})
    Poste: {employee[11] or 'N/A'}
    Email: {employee[12] or 'N/A'}
    Téléphone: {employee[13] or 'N/A'}
    Type Contrat: {employee[14] or 'N/A'}
    Date Début: {format_date_for_display(employee[15])}
    Date Fin: {format_date_for_display(employee[16])}
    Salaire Base: {employee[17] or 'N/A'} TND
    Prime: {employee[18] or 'N/A'} TND
    Type Salaire: {employee[19] or 'N/A'}"""

                self.info_text.config(state=tk.NORMAL)
                self.info_text.delete(1.0, tk.END)
                self.info_text.insert(tk.END, info_text)
                self.info_text.config(state=tk.DISABLED)

                # Activer les boutons
                self.generate_contract_btn.config(state=tk.NORMAL)
                self.edit_btn.config(state=tk.NORMAL)
                self.delete_btn.config(state=tk.NORMAL)

                self.status_var.set(f"Employé trouvé: {employee[1]} {employee[2]}")
            else:
                self.clear_search()
                Messagebox.show_info("Aucun employé trouvé avec ce matricule", "Information")

        except Exception as e:
            self.clear_search()
            print(f"Erreur lors de la recherche: {str(e)}", "Erreur")
            self.status_var.set("Erreur de recherche")

    def clear_search(self):
        self.info_text.delete(1.0, tk.END)
        self.current_employee = None
        self.generate_contract_btn.config(state=DISABLED)
        self.status_var.set("Aucun résultat")

    def check_employee_data(self, matricule):
        """Version corrigée avec gestion robuste des dates"""
        try:
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT * FROM employees WHERE matricule = %s", (matricule,))
                employee = cursor.fetchone()

                if not employee:
                    return {"status": "error", "message": "Employé non trouvé"}

                # Conversion en dictionnaire si tuple
                if isinstance(employee, tuple):
                    cols = [col[0] for col in cursor.description]
                    employee = dict(zip(cols, employee))

                # Vérification des champs obligatoires
                required_fields = [
                    'matricule', 'nom', 'prenom', 'genre',
                    'type_contrat', 'date_debut', 'salaire_base'
                ]
                missing = [f for f in required_fields if not employee.get(f)]

                # Vérification des formats de date
                date_errors = []
                date_fields = {
                    'date_naissance': "%d/%m/%Y",
                    'date_cin': "%d/%m/%Y",
                    'date_debut': "%d/%m/%Y",
                    'date_fin': "%d/%m/%Y"
                }

                for field, fmt in date_fields.items():
                    if field in employee and employee[field]:
                        try:
                            datetime.datetime.strptime(str(employee[field]), fmt)
                        except ValueError:
                            date_errors.append(field)

                return {
                    "status": "success",
                    "data": employee,
                    "missing_fields": missing,
                    "date_errors": date_errors  # Toujours retourner cette clé
                }

        except Exception as e:
            return {"status": "error", "message": str(e), "missing_fields": [], "date_errors": []}

    def diagnose_database(self):
        """Génère un rapport complet sur l'état des données"""
        try:
            with self.conn.cursor() as cursor:
                # 1. Statistiques de base
                cursor.execute("SELECT COUNT(*) FROM employees")
                total_employees = cursor.fetchone()[0]

                cursor.execute("""
                               SELECT COUNT(*)                                              as total,
                                      SUM(CASE WHEN type_contrat = 'CDD' THEN 1 ELSE 0 END) as cdd_count,
                                      SUM(CASE WHEN type_contrat = 'CDI' THEN 1 ELSE 0 END) as cdi_count,
                                      SUM(CASE WHEN date_fin IS NULL THEN 1 ELSE 0 END)     as no_end_date
                               FROM employees
                               """)
                stats = cursor.fetchone()

                # 2. Problèmes courants
                cursor.execute("""
                               SELECT COUNT(*)                as missing_essential,
                                      GROUP_CONCAT(matricule) as samples
                               FROM employees
                               WHERE matricule IS NULL
                                  OR nom IS NULL
                                  OR prenom IS NULL
                                  OR type_contrat IS NULL
                                  OR date_debut IS NULL LIMIT 5
                               """)
                problems = cursor.fetchone()

                return {
                    "total_employees": total_employees,
                    "cdd_count": stats[1],
                    "cdi_count": stats[2],
                    "no_end_date": stats[3],
                    "missing_essential": problems[0],
                    "problem_samples": problems[1] or "Aucun"
                }

        except Exception as e:
            return {"error": str(e)}

    def check_expiration_date(self):
        """
        Vérifie si la licence de l'application est toujours valide.
        Retourne True si valide, False si expirée.
        """
        try:
            # Date d'expiration fixée au 1er juillet 2025 à 09:02 (UTC)
            expiration_date = datetime.datetime(2026, 4, 1, 9, 2, tzinfo=datetime.timezone.utc)
            current_date = datetime.datetime.now(datetime.timezone.utc)

            if current_date > expiration_date:
                print(
                    "La licence de cette application a expiré.\n\n"
                    "Veuillez contacter l'administrateur ou le support technique "
                    "pour obtenir une nouvelle licence.",
                    "licence expiré.",
                    parent=self.root
                )
                return False

            # Calcul et affichage du temps restant (optionnel)
            time_left = expiration_date - current_date
            if time_left.days < 30:  # Avertir si moins d'un mois reste
                self.status_var.set(
                    f"Attention: Licence expire dans {time_left.days} jours"
                )

            return True

        except Exception as e:
            # En cas d'erreur, on considère que la licence est valide
            # mais on log l'erreur pour investigation
            logging.error(f"Erreur vérification licence: {str(e)}")
            return True


if __name__ == "__main__":
    try:
        root = ttk.Window()
        app = ContratApplication(root)

        if not app.check_expiration_date():
            root.destroy()
        else:
            root.mainloop()

    except Exception as e:
        logging.critical(f"Erreur critique: {str(e)}")
        # Correction de l'appel print() - suppression du paramètre 'parent'
        error_message = f"Erreur Initialisation: Impossible de démarrer l'application:\n{str(e)}"
        print(error_message)

        # Optionnel: Afficher aussi une messagebox d'erreur
        try:
            import tkinter.messagebox as messagebox

            messagebox.showerror("Erreur Initialisation",
                                 f"Impossible de démarrer l'application:\n{str(e)}")
        except:
            pass
